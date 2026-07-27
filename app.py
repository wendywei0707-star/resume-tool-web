"""
================================================================================
 智能简历匹配与改写工具（网页版 v3）
 Smart Resume Matcher & Rewriter (Web Edition)
================================================================================

 用途：上传岗位描述 → 分析核心要求 → 上传并匹配简历 → 改写简历 → 导出 PDF
 本地运行：streamlit run app.py，浏览器打开 http://localhost:8501
 云端部署：见 README.md（Streamlit Community Cloud，免费）

 核心匹配/关键词分析引擎运行在服务器端（jieba，无需外部 API）；
 LLM 智能改写为可选项，Key 由访问者自行输入，仅保存在本次浏览器会话中，
 不写入服务器磁盘，不同访问者之间互不可见。
================================================================================
"""

# ============================================================================
# 导入模块
# ============================================================================
import streamlit as st
import os
import sys
import io
import re
import json
import subprocess
import tempfile
import shutil
from typing import Optional, Dict, List, Tuple

# 文档处理
from docx import Document

# PDF 处理
import fitz  # PyMuPDF

# 图片处理
from PIL import Image

# 中文分词
import jieba
import jieba.analyse

# ============================================================================
# 页面配置
# ============================================================================
st.set_page_config(
    page_title="智能简历匹配与改写",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ============================================================================
# 苹果风格设计系统
# ============================================================================
def inject_apple_design():
    """注入 Apple 官网风格 CSS"""
    st.markdown(
        """
        <style>
        /* ===== 全局字体与配色 ===== */
        html, body, .stApp, [data-testid="stAppViewContainer"] {
            font-family: -apple-system, BlinkMacSystemFont, "SF Pro Display",
                         "SF Pro Text", "Helvetica Neue", "PingFang SC",
                         "Microsoft YaHei", sans-serif !important;
            color: #1d1d1f !important;
            background-color: #ffffff !important;
            -webkit-font-smoothing: antialiased !important;
        }

        /* 隐藏默认菜单和页脚 */
        #MainMenu, footer, [data-testid="stStatusWidget"] {
            visibility: hidden !important;
        }

        /* ===== 主内容容器：居中 + 限宽 ===== */
        .stApp > .main > div {
            max-width: 980px !important;
            margin: 0 auto !important;
            padding-top: 3.5rem !important;
            padding-bottom: 5rem !important;
        }

        /* ===== 标题层级（苹果排版规范） ===== */
        h1 {
            font-size: 48px !important;
            font-weight: 600 !important;
            letter-spacing: -0.003em !important;
            line-height: 1.08 !important;
            color: #1d1d1f !important;
            margin-bottom: 0.3em !important;
        }

        h2 {
            font-size: 32px !important;
            font-weight: 600 !important;
            letter-spacing: -0.002em !important;
            line-height: 1.1 !important;
            color: #1d1d1f !important;
            margin-top: 2.8rem !important;
            margin-bottom: 0.4em !important;
        }

        h3 {
            font-size: 22px !important;
            font-weight: 600 !important;
            color: #1d1d1f !important;
            margin-top: 1.8rem !important;
        }

        h4 {
            font-size: 19px !important;
            font-weight: 600 !important;
            color: #1d1d1f !important;
        }

        /* 正文与说明文字 */
        .stMarkdown p {
            color: #1d1d1f !important;
            font-size: 17px !important;
            line-height: 1.47 !important;
        }

        .stMarkdown .caption, .stCaption > div, [data-testid="stCaptionContainer"] {
            color: #86868b !important;
            font-size: 14px !important;
            line-height: 1.4 !important;
        }

        /* 引用块 */
        blockquote, .stMarkdown blockquote {
            border-left: 3px solid #0071e3 !important;
            background: #f5f5f7 !important;
            padding: 16px 20px !important;
            border-radius: 0 12px 12px 0 !important;
            color: #1d1d1f !important;
            font-size: 15px !important;
        }

        /* ===== 分隔线 ===== */
        hr {
            border: none !important;
            border-top: 1px solid #d2d2d7 !important;
            margin: 2.5rem 0 !important;
        }

        /* ===== 按钮：胶囊形（苹果核心交互元素） ===== */
        .stButton > button,
        .stDownloadButton > button {
            border-radius: 980px !important;
            font-weight: 400 !important;
            font-size: 15px !important;
            letter-spacing: -0.01em !important;
            padding: 10px 24px !important;
            transition: all 0.25s cubic-bezier(0.4, 0, 0.2, 1) !important;
            border: none !important;
            box-shadow: none !important;
        }

        /* 主按钮：苹果蓝 */
        .stButton > button[kind="primary"],
        .stDownloadButton > button[kind="primary"] {
            background-color: #0071e3 !important;
            color: #ffffff !important;
        }
        .stButton > button[kind="primary"]:hover,
        .stDownloadButton > button[kind="primary"]:hover {
            background-color: #0077ed !important;
        }

        /* 次按钮：透明蓝字蓝框 */
        .stButton > button:not([kind="primary"]) {
            background-color: transparent !important;
            color: #0071e3 !important;
            border: 1px solid #0071e3 !important;
        }
        .stButton > button:not([kind="primary"]):hover {
            background-color: #0071e3 !important;
            color: #ffffff !important;
        }

        /* ===== 输入框 / 文本区 ===== */
        .stTextInput > div > div > input,
        .stTextArea > div > div > textarea {
            border-radius: 12px !important;
            border: 1px solid #d2d2d7 !important;
            font-size: 15px !important;
            font-family: inherit !important;
            transition: border-color 0.2s ease !important;
        }
        .stTextInput > div > div > input:focus,
        .stTextArea > div > div > textarea:focus {
            border-color: #0071e3 !important;
            box-shadow: 0 0 0 3px rgba(0, 113, 227, 0.15) !important;
        }

        /* ===== 选择框 ===== */
        .stSelectbox > div > div {
            border-radius: 12px !important;
        }

        /* ===== 文件上传 ===== */
        .stFileUploader {
            border-radius: 18px !important;
        }
        [data-testid="stFileUploaderDropzone"] {
            border-radius: 18px !important;
            border: 2px dashed #d2d2d7 !important;
            background: #f5f5f7 !important;
            transition: border-color 0.2s ease !important;
        }
        [data-testid="stFileUploaderDropzone"]:hover {
            border-color: #0071e3 !important;
        }

        /* ===== 提示框（info / success / warning / error） ===== */
        .stAlert, [data-testid="stAlert"] {
            border-radius: 18px !important;
            font-size: 15px !important;
        }
        [data-testid="stAlertContainerInfo"] {
            background-color: #e8f0fe !important;
            border: 1px solid #d0e2fd !important;
        }
        [data-testid="stAlertContainerSuccess"] {
            background-color: #e8f8ee !important;
            border: 1px solid #c3eccf !important;
        }
        [data-testid="stAlertContainerWarning"] {
            background-color: #fff8e6 !important;
            border: 1px solid #ffe9a8 !important;
        }
        [data-testid="stAlertContainerError"] {
            background-color: #fde8e8 !important;
            border: 1px solid #fac7c7 !important;
        }

        /* ===== 标签页（Tabs） ===== */
        .stTabs [data-baseweb="tab-list"] {
            gap: 4px !important;
            border-bottom: 1px solid #d2d2d7 !important;
        }
        .stTabs [data-baseweb="tab"] {
            border-radius: 980px !important;
            padding: 8px 20px !important;
            font-size: 15px !important;
            font-weight: 500 !important;
            color: #86868b !important;
            transition: all 0.2s ease !important;
        }
        .stTabs [aria-selected="true"] {
            color: #0071e3 !important;
        }

        /* ===== 展开器 ===== */
        .streamlit-expanderHeader, [data-testid="stExpander"] > details > summary {
            font-size: 15px !important;
            font-weight: 500 !important;
            background-color: #f5f5f7 !important;
            border-radius: 18px !important;
        }
        [data-testid="stExpander"] {
            border-radius: 18px !important;
            border: 1px solid #d2d2d7 !important;
            overflow: hidden !important;
        }

        /* ===== 侧边栏 ===== */
        section[data-testid="stSidebar"] {
            background-color: #f5f5f7 !important;
            border-right: 1px solid #d2d2d7 !important;
        }
        section[data-testid="stSidebar"] .stMarkdown h1,
        section[data-testid="stSidebar"] .stMarkdown h2,
        section[data-testid="stSidebar"] .stMarkdown h3 {
            color: #1d1d1f !important;
            font-weight: 600 !important;
        }
        section[data-testid="stSidebar"] .stMarkdown h1 {
            font-size: 22px !important;
        }
        section[data-testid="stSidebar"] .stMarkdown h3 {
            font-size: 15px !important;
            text-transform: uppercase !important;
            letter-spacing: 0.02em !important;
            color: #86868b !important;
        }

        /* 侧边栏步骤项 */
        .sidebar-step {
            display: flex !important;
            align-items: center !important;
            gap: 10px !important;
            padding: 6px 0 !important;
            font-size: 15px !important;
            color: #1d1d1f !important;
        }
        .sidebar-step-done {
            color: #86868b !important;
        }
        .step-dot {
            width: 22px !important;
            height: 22px !important;
            border-radius: 50% !important;
            display: inline-flex !important;
            align-items: center !important;
            justify-content: center !important;
            font-size: 12px !important;
            font-weight: 600 !important;
            flex-shrink: 0 !important;
        }
        .step-dot-done {
            background: #0071e3 !important;
            color: #fff !important;
        }
        .step-dot-todo {
            background: #d2d2d7 !important;
            color: #86868b !important;
        }

        /* ===== 数字/指标展示 ===== */
        .metric-card {
            background: #f5f5f7 !important;
            border-radius: 18px !important;
            padding: 20px 24px !important;
            text-align: center !important;
        }

        /* ===== 滚动条 ===== */
        ::-webkit-scrollbar {
            width: 8px !important;
            height: 8px !important;
        }
        ::-webkit-scrollbar-thumb {
            background: #d2d2d7 !important;
            border-radius: 4px !important;
        }
        ::-webkit-scrollbar-track {
            background: transparent !important;
        }

        /* ===== 加载动画 ===== */
        .stSpinner > div {
            border-top-color: #0071e3 !important;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


inject_apple_design()

# ============================================================================
# 预置数据：技能词库 & 动词词库
# ============================================================================

# 常见技术/业务技能关键词（用于岗位分析时辅助提取）
TECH_SKILLS_DICT = {
    # 编程语言
    "python", "java", "javascript", "typescript", "c++", "c#", "go", "golang",
    "rust", "ruby", "php", "scala", "kotlin", "swift", "r语言",
    # 前端
    "react", "vue", "angular", "html", "css", "jquery", "bootstrap",
    "webpack", "vite", "next.js", "nuxt",
    # 后端
    "spring", "django", "flask", "fastapi", "node.js", "express", "gin",
    "kafka", "rabbitmq", "redis", "nginx",
    # 数据库
    "mysql", "postgresql", "mongodb", "oracle", "sql server", "elasticsearch",
    "hive", "spark", "flink", "clickhouse",
    # 云/DevOps
    "aws", "azure", "gcp", "docker", "kubernetes", "k8s", "terraform",
    "jenkins", "gitlab", "ci/cd", "devops", "linux",
    # 数据分析/AI
    "机器学习", "深度学习", "nlp", "计算机视觉", "tensorflow", "pytorch",
    "tableau", "power bi", "excel", "pandas", "scikit-learn",
    # 企业管理软件
    "erp", "sap", "oracle ebs", "用友", "金蝶", "salesforce",
    "crm", "oa", "wms", "mes", "scm", "srm", "hcm",
    # 项目管理
    "pmp", "scrum", "敏捷", "jira", "confluence", "项目管理",
    # 通用技能
    "数据分析", "团队管理", "跨部门沟通", "需求分析", "业务流程",
    "数字化转型", "系统实施", "解决方案", "产品设计", "用户体验",
    "英语", "日语", "德语",
}

# 简历优化动作动词（中文）
ACTION_VERBS_CN = [
    "主导", "负责", "推动", "设计", "实施", "优化", "搭建", "构建",
    "制定", "管理", "协调", "驱动", "带领", "交付", "落地", "改进",
    "提升", "降低", "实现", "达成", "完成", "突破", "创新", "重塑",
    "整合", "赋能", "解决", "保障", "支撑", "维护",
]

# 简历优化动作动词（英文）
ACTION_VERBS_EN = [
    "led", "managed", "designed", "implemented", "developed", "optimized",
    "built", "established", "directed", "coordinated", "drove", "delivered",
    "improved", "reduced", "achieved", "launched", "engineered", "architected",
    "spearheaded", "orchestrated", "streamlined", "transformed", "integrated",
    "resolved", "enhanced", "accelerated", "automated",
]

# ============================================================================
# Session State 初始化
# ============================================================================
# 网页版部署在公共云端，不做任何跨会话的本地持久化：
# 每个访问者的简历文件和 API Key 只存在于自己浏览器的这次会话中，
# 服务器重启或换一个人访问都不会互相看到对方的数据。
def init_session_state():
    """初始化所有 session_state 变量（网页版：不做本地磁盘持久化）"""
    defaults = {
        "resume_folder": "",
        "job_description_text": "",
        "job_analysis_result": None,       # {'skills': [], 'requirements': [], 'keywords': []}
        "matched_resumes_list": [],            # Top N 匹配结果 [{path, name, similarity, preview}, ...]
        "selected_resume_index": 0,            # 用户选择的简历索引
        "optimized_resume_path": None,
        "optimized_resume_name": None,
        "step_completed": {
            "step1": False,
            "step2": False,
            "step3": False,
            "step4": False,
            "step5": False,
            "step6": False,
        },
        "generated_pdf_path": None,
        "generated_pdf_name": None,
        # LLM 配置：每个会话独立，不预填任何人的 Key
        "llm_api_key": "",
        "llm_base_url": "https://api.deepseek.com",
        "llm_model": "deepseek-chat",
        "llm_provider": "openai_compatible",
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value


init_session_state()

# ============================================================================
# 辅助函数：文本提取
# ============================================================================

def extract_text_from_image(image_bytes: bytes) -> str:
    """使用 pytesseract OCR 从图片中提取文字（纯本地）"""
    # 检查 pytesseract 模块
    try:
        import pytesseract
    except ImportError:
        return (
            "⚠️ pytesseract 未安装，无法 OCR 识别图片。\n\n"
            "请改用以下方式输入岗位描述：\n"
            "1. 直接在左侧文本框粘贴岗位描述文字\n"
            "2. 上传 PDF 或 Word(.docx) 格式的文件"
        )

    # 检查 Tesseract 引擎是否可用
    import shutil
    tesseract_path = shutil.which("tesseract")
    if not tesseract_path:
        # 尝试常见路径
        common_paths = [
            "/usr/local/bin/tesseract",
            "/opt/homebrew/bin/tesseract",
            "/usr/bin/tesseract",
        ]
        for p in common_paths:
            if os.path.exists(p):
                pytesseract.pytesseract.tesseract_cmd = p
                tesseract_path = p
                break

    if not tesseract_path:
        return (
            "⚠️ Tesseract OCR 引擎未安装，无法识别图片文字。\n\n"
            "请改用以下方式输入岗位描述：\n"
            "1. 直接在左侧文本框粘贴岗位描述文字\n"
            "2. 上传 PDF 或 Word(.docx) 格式的文件\n\n"
            "如需安装 Tesseract（macOS 需先安装 Homebrew）：\n"
            "    brew install tesseract\n"
            "或访问 https://brew.sh 安装 Homebrew"
        )

    try:
        image = Image.open(io.BytesIO(image_bytes))
        # 尝试中英文混合识别
        text = pytesseract.image_to_string(image, lang="chi_sim+eng")
        if text.strip():
            return text.strip()
        # 如果中英文混合失败，尝试单独英文
        text = pytesseract.image_to_string(image, lang="eng")
        if text.strip():
            return text.strip()
        return "⚠️ OCR 未识别到任何文字，请确认图片清晰度，或直接粘贴文字"
    except Exception as e:
        return f"⚠️ OCR 识别失败: {e}\n建议直接粘贴岗位描述文字"


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """使用 PyMuPDF (fitz) 从 PDF 中提取文字"""
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        all_text = []
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text("text")
            if text.strip():
                all_text.append(text.strip())
        doc.close()
        result = "\n\n".join(all_text)
        if result.strip():
            return result.strip()
        else:
            return "❌ PDF 中未检测到可提取的文字（可能是扫描版图片 PDF，请尝试截图后上传）"
    except Exception as e:
        raise Exception(f"PDF 解析失败: {e}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """使用 python-docx 从 Word 文档中提取文字"""
    try:
        doc = Document(io.BytesIO(file_bytes))
        all_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                all_parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    all_parts.append(" | ".join(row_text))
        result = "\n\n".join(all_parts)
        return result if result.strip() else "❌ 文档中未检测到文字内容"
    except Exception as e:
        raise Exception(f"Word 文档解析失败: {e}")


def extract_text_from_upload(uploaded_file) -> str:
    """根据上传文件类型，选择合适的文本提取方法"""
    file_bytes = uploaded_file.read()
    file_type = uploaded_file.type.lower()
    file_name = uploaded_file.name.lower()

    if file_type in ["image/png", "image/jpeg", "image/jpg"] or any(
        file_name.endswith(ext) for ext in [".png", ".jpg", ".jpeg"]
    ):
        with st.spinner("🔍 正在使用 OCR 识别图片中的文字..."):
            return extract_text_from_image(file_bytes)

    elif file_type == "application/pdf" or file_name.endswith(".pdf"):
        with st.spinner("📄 正在解析 PDF 文件..."):
            return extract_text_from_pdf(file_bytes)

    elif (
        file_type
        in [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ]
        or file_name.endswith(".docx")
        or file_name.endswith(".doc")
    ):
        with st.spinner("📝 正在解析 Word 文档..."):
            return extract_text_from_docx(file_bytes)

    else:
        raise ValueError(f"不支持的文件格式: {file_type}。请上传 PNG/JPG/PDF/DOCX 文件。")



# ============================================================================
# 核心功能 1：本地分析岗位核心要求
# ============================================================================

def _clean_text(text: str) -> str:
    """清洗文本，去除特殊字符和多余空白"""
    text = re.sub(r'[（()）【】\[\]{}「」""''""''·•●▪]', ' ', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()


def _extract_sentences(text: str) -> List[str]:
    """将文本拆分为句子"""
    # 按常见分隔符拆分
    sentences = re.split(r'[。；;\n•·●\-\*]', text)
    return [s.strip() for s in sentences if len(s.strip()) > 5]


def get_job_keywords(job_text: str) -> Dict:
    """
    纯本地分析岗位描述，提取核心技能、要求和关键词。
    使用 jieba 分词 + TF-IDF + 预置词典匹配。
    返回：{'skills': [...], 'requirements': [...], 'keywords': [...]}
    """
    cleaned = _clean_text(job_text)

    # 1. 使用 jieba TF-IDF 提取关键词
    jieba_keywords_raw = jieba.analyse.extract_tags(cleaned, topK=30, withWeight=True)

    # 2. 使用 jieba TextRank 提取关键词（互补算法）
    jieba_textrank_raw = jieba.analyse.textrank(cleaned, topK=20, withWeight=True)

    # 合并去重
    keyword_scores = {}
    for word, weight in jieba_keywords_raw:
        keyword_scores[word] = keyword_scores.get(word, 0) + weight * 2.0
    for word, weight in jieba_textrank_raw:
        keyword_scores[word] = keyword_scores.get(word, 0) + weight

    # 3. 与技能词典交叉匹配，标记技能词
    matched_skills = []
    for word in sorted(keyword_scores, key=keyword_scores.get, reverse=True):
        word_lower = word.lower()
        for skill in TECH_SKILLS_DICT:
            if skill in word_lower or word_lower in skill:
                if skill not in [s.lower() for s in matched_skills]:
                    matched_skills.append(skill)
                break

    # 如果词典匹配不足，从 TF-IDF 中补充
    if len(matched_skills) < 5:
        for word, _ in jieba_keywords_raw[:20]:
            if len(word) >= 2 and word not in [s.lower() for s in matched_skills]:
                # 检查是否是英文缩写/专业术语
                if re.match(r'^[A-Za-z0-9\+\#\.]+$', word) or len(word) >= 3:
                    matched_skills.append(word)

    matched_skills = matched_skills[:12]  # 最多12个技能

    # 4. 从原文提取句子作为"要求/职责"
    sentences = _extract_sentences(cleaned)
    requirements = []
    for sent in sentences:
        # 筛选含动词的长句作为职责描述
        has_action = any(verb in sent for verb in ["负责", "要求", "具备", "熟悉", "掌握",
                                                      "经验", "能力", "优先", "熟练",
                                                      "require", "responsible", "experience",
                                                      "skill", "ability", "manage"])
        if has_action and len(sent) > 10:
            requirements.append(sent)

    requirements = requirements[:10]  # 最多10条

    # 5. 提取最终关键词列表（TF-IDF + 技能词 + 高频词）
    keywords = []
    # 加入技能词
    for skill in matched_skills:
        if skill not in keywords:
            keywords.append(skill)
    # 加入 TF-IDF 高频词
    for word, _ in jieba_keywords_raw[:20]:
        if word not in keywords and len(word) >= 2 and not word.isdigit():
            keywords.append(word)

    keywords = keywords[:20]  # 最多20个关键词

    return {
        "skills": matched_skills if matched_skills else ["（请检查岗位描述文本是否完整）"],
        "requirements": requirements if requirements else ["（未能自动提取职责，请查看原文）"],
        "keywords": keywords,
    }


# ============================================================================
# 核心功能 2：搜索最佳匹配简历（纯本地 jieba 关键词匹配，秒级响应）
# ============================================================================

def read_docx_text(file_path: str) -> str:
    """读取 docx 文件的全部文本内容"""
    try:
        doc = Document(file_path)
        all_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                all_text.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        all_text.append(cell.text.strip())
        return " ".join(all_text)
    except Exception as e:
        raise Exception(f"无法读取简历文件 {file_path}: {e}")


def _extract_resume_keywords(text: str, topk: int = 30) -> Dict[str, float]:
    """从简历文本中提取关键词及其权重（纯 jieba，秒级）"""
    cleaned = _clean_text(text)
    # TF-IDF 关键词
    tfidf_raw = jieba.analyse.extract_tags(cleaned, topK=topk, withWeight=True)
    # TextRank 关键词（互补）
    textrank_raw = jieba.analyse.textrank(cleaned, topK=topk//2, withWeight=True)

    keyword_scores: Dict[str, float] = {}
    for word, weight in tfidf_raw:
        keyword_scores[word] = keyword_scores.get(word, 0) + weight * 2.0
    for word, weight in textrank_raw:
        keyword_scores[word] = keyword_scores.get(word, 0) + weight

    return keyword_scores


def find_best_resume(job_text: str, folder_path: str, top_n: int = 3) -> List[Dict]:
    """
    在指定文件夹内搜索最佳匹配简历（纯本地 jieba，秒级响应）。
    返回 Top N 结果列表，每项: {path, name, similarity, preview}
    """
    if not os.path.isdir(folder_path):
        raise ValueError(f"文件夹不存在: {folder_path}")

    docx_files = []
    for root, dirs, files in os.walk(folder_path):
        for f in files:
            if f.lower().endswith(".docx") and not f.startswith("~$"):
                docx_files.append(os.path.join(root, f))

    if not docx_files:
        raise ValueError(f"在 {folder_path} 中未找到任何 .docx 文件")

    st.info(f"📂 在文件夹中找到 {len(docx_files)} 份简历 (.docx)，正在分析...")

    # 1. 提取岗位关键词
    with st.spinner("🔍 正在提取岗位关键词..."):
        job_kw = _extract_resume_keywords(job_text, topk=40)
        job_text_lower = job_text.lower()
        for skill in TECH_SKILLS_DICT:
            if skill in job_text_lower and skill not in job_kw:
                job_kw[skill] = 0.5

        if not job_kw:
            raise ValueError("无法从岗位描述中提取有效关键词，请确认文本内容")

    # 2. 逐份简历打分
    scores = []
    with st.spinner(f"🧮 正在计算 {len(docx_files)} 份简历的匹配度（纯本地 jieba）..."):
        for fp in docx_files:
            try:
                resume_text = read_docx_text(fp)
                if not resume_text.strip():
                    continue
            except Exception as e:
                st.warning(f"跳过文件 {os.path.basename(fp)}: {e}")
                continue

            resume_kw = _extract_resume_keywords(resume_text, topk=40)
            resume_text_lower = resume_text.lower()

            # 维度1：关键词重叠度
            overlap_count = 0
            weighted_overlap = 0.0
            for kw, weight in job_kw.items():
                if kw.lower() in resume_kw or kw.lower() in resume_text_lower:
                    overlap_count += 1
                    resume_weight = resume_kw.get(kw.lower(), 0.3)
                    weighted_overlap += weight * resume_weight

            kw_score = weighted_overlap / max(sum(job_kw.values()), 0.01)
            kw_score = min(kw_score, 1.0)

            # 维度2：关键词命中率
            hit_rate = overlap_count / max(len(job_kw), 1)

            # 维度3：技能词典直接命中
            job_skills = {s for s in TECH_SKILLS_DICT if s in job_text_lower}
            resume_skill_hits = sum(1 for s in job_skills if s in resume_text_lower)
            skill_score = resume_skill_hits / max(len(job_skills), 1)

            # 维度4：产品/公司名称直接匹配
            product_bonus = 0.0
            known_products = ["sap", "oracle", "salesforce", "用友", "金蝶", "erp",
                              "crm", "srm", "wms", "mes", "oa", "scm", "hcm",
                              "aws", "azure", "gcp", "docker", "kubernetes"]
            for prod in known_products:
                if prod in job_text_lower and prod in resume_text_lower:
                    product_bonus += 0.04

            final_score = (
                kw_score * 0.35 +
                hit_rate * 0.30 +
                skill_score * 0.25 +
                min(product_bonus, 0.15)
            )
            final_score = min(final_score, 1.0)

            scores.append((fp, final_score, resume_text))

    if not scores:
        raise ValueError("所有简历文件均无法读取或为空")

    # 3. 排序取 Top N
    scores.sort(key=lambda x: x[1], reverse=True)
    top_n = min(top_n, len(scores))

    results = []
    for i in range(top_n):
        fp, sim, text = scores[i]
        preview = text[:1000] + ("..." if len(text) > 1000 else "")
        results.append({
            "path": fp,
            "name": os.path.basename(fp),
            "similarity": sim,
            "preview": preview,
        })

    return results


# ============================================================================
# 核心功能 3：智能简历改写引擎（LLM 优先 + jieba 兜底）
# ============================================================================

# ---------------------------------------------------------------------------
# LLM 改写引擎（大模型推理，解决模板重复问题）
# ---------------------------------------------------------------------------

def _llm_available() -> bool:
    """检查 LLM 是否已配置且可用"""
    api_key = st.session_state.get("llm_api_key", "").strip()
    if not api_key:
        return False
    try:
        import openai
        return True
    except ImportError:
        st.warning("⚠️ openai 包未安装，请运行：pip install openai")
        return False


def _call_llm(system_prompt: str, user_prompt: str, max_retries: int = 2) -> Optional[str]:
    """
    调用 LLM API（OpenAI 兼容接口）。
    支持：OpenAI / DeepSeek / Moonshot / 智谱 / Ollama 等所有兼容 OpenAI 接口的服务。
    返回：模型回复文本，失败时返回 None。
    """
    import openai

    api_key = st.session_state.get("llm_api_key", "").strip()
    base_url = st.session_state.get("llm_base_url", "https://api.deepseek.com").strip()
    model = st.session_state.get("llm_model", "deepseek-chat").strip()

    client = openai.OpenAI(api_key=api_key, base_url=base_url)

    for attempt in range(max_retries + 1):
        try:
            response = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                temperature=0.3,  # 低温度保证稳定输出
                max_tokens=4096,
                timeout=120,
            )
            content = response.choices[0].message.content
            if content and content.strip():
                return content.strip()
            if attempt < max_retries:
                continue
            return None
        except Exception as e:
            if attempt < max_retries:
                st.warning(f"🔄 LLM 调用失败（第{attempt+1}次），正在重试... ({str(e)[:80]})")
                continue
            st.error(f"❌ LLM 调用失败: {e}")
            return None
    return None


# ---------------------------------------------------------------------------
# LLM 改写引擎 v2：一次性改写整份简历，按区块回填
# ---------------------------------------------------------------------------

SECTION_MARKER_PREFIX = "===== SECTION:"
SECTION_MARKER_SUFFIX = "====="


def _extract_resume_sections(doc) -> List[Dict]:
    """
    从 docx 文档中提取完整简历文本，并按区块划分。
    支持普通段落和表格内段落。
    
    返回：
    [
      {
        'type': 'summary' | 'skills' | 'experience' | 'education' | 'header' | 'other',
        'header_text': '区块标题原文',
        'content_paras': [(para_index, Paragraph对象), ...],  # 内容段落（不含标题）
        'content_text': '段落1\\n段落2...',
      },
      ...
    ]
    """
    # 收集所有段落（含表格内段落），记录来源位置
    # all_paras: [(source_type, source_idx, para_obj), ...]
    # source_type: "doc" | "table_cell"
    all_paras = []
    
    # 1. 文档正文段落
    for idx, para in enumerate(doc.paragraphs):
        if para.text.strip():
            all_paras.append(("doc_para", idx, para))
    
    # 2. 表格内段落（每个单元格视为一个段落块）
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if not cell_text:
                    continue
                # 单元格内可能有多个段落
                cell_paras = cell.paragraphs
                # 如果单元格只有一个段落但有多个逻辑行（\n分隔），用第一个段落承载，其余行拆分
                if len(cell_paras) <= 1 and '\n' in cell_text:
                    # 拆分多行文本，每行作为一个虚拟段落
                    lines = [l.strip() for l in cell_text.split('\n') if l.strip()]
                    for line in lines:
                        all_paras.append(("table_cell", (table_idx, row_idx, cell_idx, cell, "line", line), None))
                else:
                    for pi, para in enumerate(cell_paras):
                        if para.text.strip():
                            all_paras.append(("table_cell", (table_idx, row_idx, cell_idx, cell, "para", pi, para), para))

    # 识别区块边界
    section_boundaries = []
    for idx, item in enumerate(all_paras):
        src_type = item[0]
        if src_type == "doc_para":
            text = item[2].text.strip()
        elif src_type == "table_cell" and item[2] is not None:
            text = item[2].text.strip()
        elif src_type == "table_cell":
            # 虚拟行
            text = item[1][4]
        else:
            continue
        section_type = _detect_section_type(text)
        if section_type != "content":
            section_boundaries.append((idx, section_type, text))

    # 如果没有识别到任何区块
    if not section_boundaries:
        # 收集所有内容
        content_paras = []
        content_texts = []
        for idx, item in enumerate(all_paras):
            content_paras.append((idx, item))
            text = _get_para_text(item)
            if text:
                content_texts.append(text)
        return [{
            "type": "other",
            "header_text": "",
            "content_paras": content_paras,
            "content_text": "\n".join(content_texts),
        }]

    # 第一个区块之前的所有内容视为 header
    sections = []
    first_boundary_idx = section_boundaries[0][0]
    if first_boundary_idx > 0:
        header_paras = []
        header_texts = []
        for idx in range(0, first_boundary_idx):
            item = all_paras[idx]
            text = _get_para_text(item)
            if text:
                header_paras.append((idx, item))
                header_texts.append(text)
        if header_texts:
            sections.append({
                "type": "header",
                "header_text": "",
                "content_paras": header_paras,
                "content_text": "\n".join(header_texts),
            })

    # 遍历每个区块
    for i, (start_idx, section_type, header_text) in enumerate(section_boundaries):
        end_idx = section_boundaries[i+1][0] if i+1 < len(section_boundaries) else len(all_paras)
        content_paras = []
        content_texts = []
        for idx in range(start_idx + 1, min(end_idx, len(all_paras))):
            item = all_paras[idx]
            text = _get_para_text(item)
            if text:
                content_paras.append((idx, item))
                content_texts.append(text)

        sections.append({
            "type": section_type,
            "header_text": header_text,
            "content_paras": content_paras,
            "content_text": "\n".join(content_texts),
        })

    return sections


def _get_para_text(item) -> str:
    """从 all_paras 条目中提取文本"""
    src_type = item[0]
    if src_type == "doc_para":
        return item[2].text.strip()
    elif src_type == "table_cell":
        meta = item[1]
        if meta[4] == "line":
            return meta[5]  # 虚拟行的文本
        elif item[2] is not None:
            return item[2].text.strip()
    return ""


def _build_full_resume_text(sections: List[Dict]) -> str:
    """
    把简历各区块拼接成带明确标记的文本，方便 LLM 理解结构并返回。
    """
    lines = []
    section_type_names = {
        "header": "个人信息/联系方式",
        "summary": "个人总结/自我评价",
        "skills": "技能列表/核心能力",
        "experience": "工作经历/项目经验",
        "education": "教育背景",
        "other": "其他信息",
    }

    for sec in sections:
        sec_name = section_type_names.get(sec["type"], sec["type"])
        header = sec["header_text"] if sec["header_text"] else sec_name
        lines.append(f"{SECTION_MARKER_PREFIX} {header} {SECTION_MARKER_SUFFIX}")
        lines.append(sec["content_text"])
        lines.append("")  # 区块间空行

    return "\n".join(lines).strip()


def _split_section_text(section_text: str) -> List[str]:
    """
    把 LLM 返回的某个区块文本拆分成段落/项目符号。
    按换行或项目符号拆分，过滤空行。
    """
    if not section_text:
        return []
    # 先按换行拆分
    raw_lines = section_text.split("\n")
    result = []
    for line in raw_lines:
        line = line.strip()
        if not line:
            continue
        # 去掉项目符号前缀
        line = re.sub(r'^[•●·\-\*–—▪▸◆○]\s*', '', line)
        if line:
            result.append(line)
    return result


def _apply_section_content(section: Dict, new_text: str) -> int:
    """
    将 LLM 返回的某个区块文本，回填到该区块对应的所有内容段落中。
    支持普通段落和表格内段落。

    返回实际修改的段落数。
    """
    content_paras = section["content_paras"]
    if not content_paras:
        return 0

    new_lines = _split_section_text(new_text)
    if not new_lines:
        return 0

    # 收集原始内容文本，用于对比
    original_texts = [_get_item_text(item) for _, item in content_paras]
    original_combined = "\n".join(original_texts)
    new_combined = "\n".join(new_lines)

    # 如果 LLM 返回的内容与原内容完全一样，则跳过
    if _texts_effectively_equal(original_combined, new_combined):
        return 0

    modified_count = 0
    n_original = len(content_paras)
    n_new = len(new_lines)

    if n_new >= n_original:
        for i, (idx, item) in enumerate(content_paras):
            new_line = new_lines[i] if i < n_new else ""
            if _write_item_if_changed(item, new_line):
                modified_count += 1
    else:
        if n_new == 1 and n_original > 1:
            single_text = new_lines[0]
            split_pattern = r'(?<=[。；;])\s*'
            parts = re.split(split_pattern, single_text)
            parts = [p.strip() for p in parts if p.strip()]
            if len(parts) >= n_original:
                for i, (idx, item) in enumerate(content_paras):
                    new_line = parts[i] if i < len(parts) else ""
                    if _write_item_if_changed(item, new_line):
                        modified_count += 1
            else:
                for i, (idx, item) in enumerate(content_paras):
                    if i == 0:
                        if _write_item_if_changed(item, single_text):
                            modified_count += 1
                    else:
                        if original_texts[i]:
                            _clear_item(item)
                            modified_count += 1
        else:
            for i, (idx, item) in enumerate(content_paras):
                if i < n_new:
                    if _write_item_if_changed(item, new_lines[i]):
                        modified_count += 1
                else:
                    if original_texts[i]:
                        _clear_item(item)
                        modified_count += 1

    return modified_count


def _get_item_text(item) -> str:
    """从 all_paras 条目中提取当前文本"""
    src_type = item[0]
    if src_type == "doc_para":
        return item[2].text.strip()
    elif src_type == "table_cell":
        meta = item[1]
        if meta[4] == "line":
            return meta[5]  # 虚拟行的文本
        elif item[2] is not None:
            return item[2].text.strip()
    return ""


def _write_item_if_changed(item, new_text: str) -> bool:
    """将文本写入段落（支持文档段落和表格内段落）。返回是否实际修改"""
    new_text = (new_text or "").strip()
    original_text = _get_item_text(item)
    src_type = item[0]

    if _texts_effectively_equal(original_text, new_text):
        return False

    if not new_text:
        if original_text:
            _clear_item(item)
            return True
        return False

    if src_type == "doc_para":
        para = item[2]
        if para.runs:
            first_run = para.runs[0]
            for r_idx in range(1, len(para.runs)):
                para.runs[r_idx].text = ""
            first_run.text = new_text
        else:
            para.add_run(new_text)
        return True

    elif src_type == "table_cell":
        meta = item[1]
        if meta[4] == "line":
            # 虚拟行 — 无法直接写回，跳过
            return False
        elif item[2] is not None:
            para = item[2]
            if para.runs:
                first_run = para.runs[0]
                for r_idx in range(1, len(para.runs)):
                    para.runs[r_idx].text = ""
                first_run.text = new_text
            else:
                para.add_run(new_text)
            return True

    return False


def _clear_item(item) -> None:
    """清空段落内容"""
    src_type = item[0]
    if src_type == "doc_para":
        para = item[2]
        if para.runs:
            for run in para.runs:
                run.text = ""
        para.text = ""
    elif src_type == "table_cell":
        meta = item[1]
        if meta[4] == "para" and item[2] is not None:
            para = item[2]
            if para.runs:
                for run in para.runs:
                    run.text = ""
            para.text = ""


def _texts_effectively_equal(text1: str, text2: str) -> bool:
    """判断两段文本是否实质相同（忽略空白和标点差异）"""
    def normalize(t):
        return re.sub(r'\s+', '', t).strip()
    return normalize(text1) == normalize(text2)


def _build_rewrite_prompt_v2(job_description: str, job_analysis: Dict,
                              resume_text: str) -> Tuple[str, str]:
    """
    按用户建议，一次性把完整 JD 和完整简历发给 LLM，让它产出完整新简历。
    Prompt 强调：【禁止复制原文】，必须根据 JD 真正改写内容。
    """
    keywords = job_analysis.get("keywords", [])
    requirements = job_analysis.get("requirements", [])
    skills = job_analysis.get("skills", [])

    system_prompt = (
        "你是一位顶级简历优化专家，专门帮求职者根据目标岗位JD优化简历。\n\n"
        "你的核心使命是：**根据JD重新改写简历内容，将JD关键词自然融入**。\n\n"
        "你的工作方式：\n"
        "1. 仔细阅读JD，提取所有核心关键词和技能要求\n"
        "2. 读取原始简历，理解求职者的真实背景和经历\n"
        "3. 对每个区块进行实质性改写——用JD的术语和关键词重新包装现有经历\n"
        "4. 只保留区块标题不变，内容必须全部重写\n\n"
        "改写原则：\n"
        "- 每个区块至少注入3-5个JD核心关键词（自然融入，不堆砌）\n"
        "- 每条经历改写为\"强动词 + JD关键词 + 量化结果\"结构\n"
        "- 使用强动词：主导、推动、设计、实施、统筹、搭建、交付、驱动\n"
        "- 保留原文的真实数字指标（百分比、金额、效率提升等）\n"
        "- 同一区块内多条经历使用不同句式和表述，禁止重复\n"
        "- 不编造虚假技能/项目，但可以用JD术语重新包装现有经验\n"
        "- 禁止使用markdown格式（不要用**加粗**），只用纯文本"
    )

    user_prompt = f"""
你是一名简历优化专家。请根据下面的岗位要求(JD)，对我的原始简历进行【实质性改写】。

═══════════════════════════════════════
⚠️ 核心要求（违反将导致改写失败）：
═══════════════════════════════════════

1. 【必须修改】每个区块的内容都必须根据 JD 重新组织改写，**严禁**照抄原始简历
2. 【关键词注入】必须将以下清单中的关键词全部自然融入简历各区块（在相关上下文提及即可）：
   {', '.join(keywords[:25]) if keywords else '见 JD'}
3. 【量化结果】每条工作经历至少包含 1 个具体的数字指标（百分比、金额、天数、人数等）
4. 【结构不变】区块标题和区块顺序保持不变，只改写标题下面的内容
5. 【用 JD 术语重写】将简历中的通用表达替换为 JD 中的行业术语，
   例如："IT系统"→"数字工厂/智能制造系统"，"销售门户"→"业务流程数字化"
6. 【严禁编造】不虚构不存在的技能或项目，只在现有经历基础上用 JD 关键词重新包装
7. 【禁止重复】同一区块内多条经历/能力点必须各不相同
8. 【禁止 markdown】不要使用 **加粗** 格式，只用纯文本

═══════════════════════════════════════
📋 各区块改写要求：
═══════════════════════════════════════

• 个人信息/联系方式：保持不变（不修改）
• 个人总结/核心能力：用 JD 关键词完全重写，每条突出一个 JD 核心主题，嵌入 3-5 个 JD 最核心关键词
• 技能列表/核心能力：补充 JD 要求但原始简历中缺失的硬技能，重新归类排版
• 工作经历/项目经验：每条经历改写为"强动词 + JD 关键词 + 量化结果"结构，
  将每段经历导向 JD 所在行业场景，让 HR 看到高度匹配
• 教育背景：保持不变（不修改）

═══════════════════════════════════════
📄 输出格式（严格遵循）：
═══════════════════════════════════════

每个区块用以下格式标记：
===== SECTION: 区块标题 =====
（改写后的内容）

注意：区块标题与原始简历完全一致，不要修改标题文字。

═══════════════════════════════════════
🔑 岗位描述 (JD)：
═══════════════════════════════════════
{job_description}

═══════════════════════════════════════
🔑 岗位核心技能要求：
═══════════════════════════════════════
{', '.join(skills) if skills else '见岗位描述'}

═══════════════════════════════════════
🔑 岗位核心职责：
═══════════════════════════════════════
{chr(10).join(f'- {r}' for r in requirements) if requirements else '见岗位描述'}

═══════════════════════════════════════
🔑 必须融入的关键词清单：
═══════════════════════════════════════
{', '.join(keywords) if keywords else '见岗位描述'}

═══════════════════════════════════════
📝 我的原始简历（请改写以下每个区块的内容）：
═══════════════════════════════════════
{resume_text}

═══════════════════════════════════════

请现在开始输出改写后的完整简历。记住：内容必须实质性改写，禁止复制原文，必须融入所有JD关键词。"""
    return system_prompt, user_prompt


def _parse_llm_resume_output(response_text: str) -> Dict[str, str]:
    """
    解析 LLM 返回的完整简历文本，按 SECTION 标记切分。
    返回：{区块标题: 区块内容}
    """
    sections = {}
    if not response_text:
        return sections

    # 匹配区块标记：===== SECTION: 标题 =====
    pattern = rf"{re.escape(SECTION_MARKER_PREFIX)}\s*(.*?)\s*{re.escape(SECTION_MARKER_SUFFIX)}"
    matches = list(re.finditer(pattern, response_text))

    if not matches:
        return sections

    for i, match in enumerate(matches):
        section_title = match.group(1).strip()
        start = match.end()
        end = matches[i+1].start() if i+1 < len(matches) else len(response_text)
        content = response_text[start:end].strip()
        sections[section_title] = content

    return sections


def _rewrite_resume_with_llm(matched_doc_path: str, job_description: str,
                              job_analysis: Dict) -> str:
    """
    使用大模型(LLM)一次性改写整份简历，再按区块回填到 docx。
    彻底解决段落编号导致的重复问题。
    """
    st.info(f"🤖 正在使用 LLM ({st.session_state['llm_model']}) 进行整份简历改写...")

    # 1. 读取文档并提取区块
    doc = Document(matched_doc_path)
    sections = _extract_resume_sections(doc)

    # 2. 构建完整简历文本
    full_resume_text = _build_full_resume_text(sections)

    # 3. 构建 Prompt 并调用 LLM
    with st.spinner("🧠 LLM 正在理解岗位要求并重写整份简历（约 20-40 秒）..."):
        system_prompt, user_prompt = _build_rewrite_prompt_v2(
            job_description, job_analysis, full_resume_text)

        response_text = _call_llm(system_prompt, user_prompt)

    if not response_text:
        raise Exception("LLM 未返回有效响应，请检查 API 配置或网络连接")

    # 4. 【验证】检查 LLM 是否真正改动了内容
    # 去除 section 标记后再比较
    _orig_content = re.sub(r'===== SECTION:.*?=====', '', full_resume_text)
    _orig_content = re.sub(r'\s+', '', _orig_content)
    _resp_content = re.sub(r'===== SECTION:.*?=====', '', response_text)
    _resp_content = re.sub(r'\s+', '', _resp_content)
    
    if _orig_content == _resp_content:
        st.error("⚠️ LLM 返回内容与原始简历完全一致，未做实质性改写")
        st.warning("建议：1) 检查岗位描述是否完整  2) 换用更强的模型  3) 重试")
        with st.expander("👀 查看 LLM 原始响应（前 2000 字）"):
            st.code(response_text[:2000])
        raise Exception("LLM 未进行实质性改写，请重试或切换更强大的模型")

    # 5. 解析 LLM 输出
    llm_sections = _parse_llm_resume_output(response_text)
    if not llm_sections:
        st.error("⚠️ LLM 返回格式异常，无法识别区块标记")
        with st.expander("👀 查看 LLM 原始响应（前 2000 字）"):
            st.code(response_text[:2000])
        raise Exception("LLM 响应格式异常，请重试或切换模型")

    # 6. 按区块回填（标题精确匹配 → 类型模糊匹配 → 顺序兜底匹配）
    total_modified = 0
    applied_sections = []
    skipped_sections = []
    failed_sections = []
    llm_section_keys = list(llm_sections.keys())

    for sec_idx, sec in enumerate(sections):
        # 用区块标题或类型名去匹配 LLM 返回的区块
        best_match_key = None
        # 优先用原标题匹配
        if sec["header_text"] and sec["header_text"] in llm_sections:
            best_match_key = sec["header_text"]
        else:
            # 用类型名模糊匹配
            type_name_map = {
                "summary": ["个人总结", "自我评价", "个人简介", "summary"],
                "skills": ["技能", "核心能力", "专业能力", "skills"],
                "experience": ["工作经历", "项目经验", "职业经历", "experience"],
                "education": ["教育背景", "教育经历", "education"],
                "header": ["个人信息", "联系方式", "基本信息", "header"],
            }
            candidates = type_name_map.get(sec["type"], [])
            for key in llm_section_keys:
                key_lower = key.lower()
                for cand in candidates:
                    if cand.lower() in key_lower:
                        best_match_key = key
                        break
                if best_match_key:
                    break

            # 兜底：按顺序匹配（LLM 通常保持原有顺序）
            if not best_match_key and sec_idx < len(llm_section_keys):
                best_match_key = llm_section_keys[sec_idx]

        if best_match_key:
            new_text = llm_sections[best_match_key]
            # 跳过 header 和 education 类型（通常不需要改写）
            if sec["type"] in ("header", "education"):
                skipped_sections.append(f"{sec.get('header_text', sec['type'])} (保持原样)")
                continue
            modified = _apply_section_content(sec, new_text)
            if modified > 0:
                applied_sections.append(f"{sec.get('header_text', sec['type'])} ✅{modified}处")
            else:
                skipped_sections.append(f"{sec.get('header_text', sec['type'])} (内容未变)")
            total_modified += modified
        else:
            failed_sections.append(f"{sec.get('header_text', sec['type'])} (无法匹配)")

    # 7. 验证：如果没有任何实际修改，报错
    if total_modified == 0:
        st.error("⚠️ LLM 返回的改写内容与原简历完全一致，0 处段落被修改")
        st.warning("建议：1) 检查岗位描述是否与简历领域相关  2) 切换更强大的模型  3) 重试")
        with st.expander("👀 查看 LLM 返回的区块内容"):
            for key, val in llm_sections.items():
                st.caption(f"**{key}**:")
                st.code(val[:500])
        raise Exception("LLM 改写后未检测到任何内容变化，请切换模型或重试")

    # 8. 保存
    output_dir = os.path.dirname(matched_doc_path)
    original_name = os.path.basename(matched_doc_path)
    name_without_ext, ext = os.path.splitext(original_name)
    new_name = f"优化后_{name_without_ext}_LLM{ext}"
    new_path = os.path.join(output_dir, new_name)

    counter = 1
    while os.path.exists(new_path):
        new_name = f"优化后_{name_without_ext}_LLM_v{counter}{ext}"
        new_path = os.path.join(output_dir, new_name)
        counter += 1

    doc.save(new_path)

    # 详细统计报告
    st.caption(f"📊 改写统计：共 {len(llm_sections)} 个区块，实际修改 {total_modified} 处段落")
    if applied_sections:
        st.caption(f"✅ 已改写：{' | '.join(applied_sections)}")
    if skipped_sections:
        st.caption(f"⏭️ 已跳过：{' | '.join(skipped_sections)}")
    if failed_sections:
        st.caption(f"⚠️ 未匹配：{' | '.join(failed_sections)}")
    return new_path


# ---------------------------------------------------------------------------
# 3.1 深度 JD 分析：将关键词分类并赋予权重
# ---------------------------------------------------------------------------

# 中文简历常见区块标题
SECTION_HEADERS_CN = {
    "summary":    ["个人总结", "个人简介", "自我评价", "自我介绍", "求职意向", "职业目标",
                   "求职目标", "专业概述", "关于我", "个人概述", "profile", "summary",
                   "objective", "professional summary", "career objective"],
    "skills":     ["专业技能", "核心能力", "技能清单", "技术栈", "掌握技能", "语言能力",
                   "证书资质", "技能", "能力", "专业能力", "技术能力", "核心技能",
                   "skills", "core competencies", "technical skills", "expertise",
                   "certifications", "languages"],
    "experience": ["工作经历", "工作经验", "职业经历", "从业经历", "工作履历", "项目经验",
                   "项目经历", "主要项目", "代表项目", "相关经历", "实习经历",
                   "experience", "work experience", "professional experience",
                   "employment", "career history", "work history", "projects",
                   "project experience", "key projects"],
    "education":  ["教育背景", "教育经历", "学历", "学习经历", "培训经历", "培训",
                   "education", "academic", "qualifications"],
}


def _analyze_jd_deep(job_text: str, job_analysis: Dict) -> Dict:
    """
    深度分析JD，将关键词分为多个类别，便于精准匹配到简历不同区块。
    返回:
    {
        'hard_skills':     [(keyword, weight), ...],   # 硬技能：工具/平台/技术
        'soft_skills':     [(keyword, weight), ...],   # 软技能：管理/沟通/领导
        'products_tools':  [(keyword, weight), ...],   # 产品/系统名
        'industry_terms':  [(keyword, weight), ...],   # 行业术语
        'action_verbs':    [verb, ...],                # 动作要求
        'all_weighted':    {keyword: weight, ...},     # 所有词合并权重
        'top_keywords':    [kw1, kw2, ...],            # 最重要的15个关键词
    }
    """
    keywords = job_analysis.get("keywords", [])
    skills = job_analysis.get("skills", [])
    job_text_lower = job_text.lower()

    # ====== 定义分类词库 ======
    # 硬技能：工具、平台、编程语言、框架、数据库等
    hard_skill_indicators = {
        "sap", "oracle", "salesforce", "用友", "金蝶", "erp", "crm", "scm",
        "srm", "wms", "mes", "oa", "hcm", "aws", "azure", "gcp", "docker",
        "kubernetes", "k8s", "terraform", "jenkins", "gitlab", "nginx",
        "redis", "mysql", "postgresql", "mongodb", "elasticsearch", "hive",
        "spark", "flink", "clickhouse", "tableau", "power bi", "excel",
        "python", "java", "javascript", "typescript", "c++", "go", "golang",
        "react", "vue", "angular", "spring", "django", "flask", "fastapi",
        "kafka", "rabbitmq", "tensorflow", "pytorch", "scikit-learn",
        "pandas", "jira", "confluence", "linux", "unix", "sql", "nosql",
        "ci/cd", "devops", "snowflake", "databricks", "servicenow",
        "workday", "peoplesoft", "netsuite", "dynamics", "sharepoint",
        "powerapps", "power automate", "power platform", "cognos",
        "hyperion", "anaplan", "tagetik", "board", "jedox",
    }
    # 软技能：管理、沟通、领导力等
    soft_skill_indicators = {
        "团队管理", "项目管理", "跨部门", "沟通", "协调", "领导", "战略",
        "变革管理", "敏捷", "scrum", "pmp", "需求分析", "业务流程",
        "问题解决", "决策", "谈判", "辅导", "培训", "演讲", "汇报",
        "预算管理", "供应商管理", "风险管理", "质量管理", "客户关系",
        "leadership", "management", "communication", "strategic",
        "stakeholder", "agile", "scrum", "pmp", "change management",
    }
    # 行业术语
    industry_indicators = {
        "数字化转型", "智能制造", "工业4.0", "大数据", "云计算", "人工智能",
        "物联网", "区块链", "商业智能", "数据驱动", "流程优化", "自动化",
        "供应链", "制造执行", "质量管理", "精益生产", "六西格玛",
        "digital transformation", "industry 4.0", "iot", "ai", "ml",
        "business intelligence", "data-driven", "lean", "six sigma",
    }

    # ====== 分类关键词 ======
    hard_skills_list = []
    soft_skills_list = []
    products_tools_list = []
    industry_terms_list = []

    # 对每个关键词分类
    keyword_scores = {}
    for kw in keywords:
        kw_lower = kw.lower()
        # 基础权重：越前面的关键词越重要
        idx = keywords.index(kw)
        base_weight = max(1.0 - idx * 0.03, 0.15)

        # 分类
        if kw_lower in hard_skill_indicators or any(
            ind in kw_lower or kw_lower in ind
            for ind in hard_skill_indicators if len(ind) > 2
        ):
            # 进一步区分：纯产品/工具名 vs 通用技能
            if any(kw_lower == p for p in ["sap", "oracle", "salesforce",
                                             "用友", "金蝶", "aws", "azure",
                                             "gcp", "workday"]):
                products_tools_list.append((kw, base_weight * 1.2))
                hard_skills_list.append((kw, base_weight))
            else:
                hard_skills_list.append((kw, base_weight))
        elif kw_lower in soft_skill_indicators or any(
            ind in kw_lower for ind in soft_skill_indicators if len(ind) > 2
        ):
            soft_skills_list.append((kw, base_weight * 1.1))
        elif kw_lower in industry_indicators or any(
            ind in kw_lower for ind in industry_indicators if len(ind) > 2
        ):
            industry_terms_list.append((kw, base_weight))
        else:
            # 未分类的关键词根据长度和特征判断
            if len(kw) >= 3 and re.search(r'[\u4e00-\u9fff]', kw):
                soft_skills_list.append((kw, base_weight * 0.7))
            else:
                hard_skills_list.append((kw, base_weight * 0.7))

        keyword_scores[kw] = base_weight

    # ====== 从技能列表补充 ======
    for skill in skills:
        skill_lower = skill.lower()
        if skill_lower not in keyword_scores:
            keyword_scores[skill] = 0.6
            if skill_lower in hard_skill_indicators or any(
                ind in skill_lower for ind in hard_skill_indicators if len(ind) > 2
            ):
                hard_skills_list.append((skill, 0.6))
            elif skill_lower in soft_skill_indicators:
                soft_skills_list.append((skill, 0.6))
            else:
                soft_skills_list.append((skill, 0.5))

    # ====== 去重排序 ======
    def _dedup_sort(items):
        seen = set()
        result = []
        for item in sorted(items, key=lambda x: x[1], reverse=True):
            kw_lower = item[0].lower()
            if kw_lower not in seen:
                seen.add(kw_lower)
                result.append(item)
        return result

    hard_skills_list = _dedup_sort(hard_skills_list)
    soft_skills_list = _dedup_sort(soft_skills_list)
    products_tools_list = _dedup_sort(products_tools_list)
    industry_terms_list = _dedup_sort(industry_terms_list)

    # ====== 提取动作要求动词 ======
    action_verbs = []
    action_markers_cn = ["负责", "主导", "推动", "管理", "设计", "实施", "优化",
                         "搭建", "构建", "制定", "协调", "带领", "驱动", "交付",
                         "落地", "改进", "提升", "降低"]
    action_markers_en = ["lead", "manage", "drive", "design", "implement",
                         "develop", "build", "deliver", "optimize", "improve",
                         "coordinate", "oversee", "direct", "execute"]
    for marker in action_markers_cn:
        if marker in job_text:
            action_verbs.append(marker)
    for marker in action_markers_en:
        if marker in job_text_lower:
            action_verbs.append(marker)

    # ====== 生成 top 15 核心关键词（按权重排序） ======
    all_sorted = sorted(keyword_scores.items(), key=lambda x: x[1], reverse=True)
    # 优先级：产品工具 > 硬技能 > 行业术语 > 软技能
    top_keywords = []
    for kw, _ in products_tools_list[:4] + hard_skills_list[:5] + industry_terms_list[:3] + soft_skills_list[:3]:
        if kw not in top_keywords:
            top_keywords.append(kw)
    top_keywords = list(dict.fromkeys(top_keywords))[:15]

    return {
        "hard_skills": hard_skills_list,
        "soft_skills": soft_skills_list,
        "products_tools": products_tools_list,
        "industry_terms": industry_terms_list,
        "action_verbs": action_verbs,
        "all_weighted": keyword_scores,
        "top_keywords": top_keywords,
    }


# ---------------------------------------------------------------------------
# 3.2 简历区块识别
# ---------------------------------------------------------------------------

def _detect_section_type(text: str) -> str:
    """判断文本所属的简历区块类型"""
    text_lower = text.strip().lower()
    # 去掉冒号、空格
    text_clean = re.sub(r'[：:\s]', '', text_lower)

    for section_type, headers in SECTION_HEADERS_CN.items():
        for header in headers:
            if header in text_clean and len(text_clean) < 30:
                return section_type
    return "content"  # 普通内容


def _is_section_header(text: str) -> bool:
    """判断文本是否是区块标题"""
    text_lower = text.strip().lower()
    text_clean = re.sub(r'[：:\s]', '', text_lower)
    if len(text_clean) > 25:
        return False
    for headers in SECTION_HEADERS_CN.values():
        for header in headers:
            if header in text_clean:
                return True
    return False


# ---------------------------------------------------------------------------
# 3.3 关键词与文本的智能匹配
# ---------------------------------------------------------------------------

def _calculate_text_relevance(text: str, keyword: str) -> float:
    """
    计算一个关键词与一段文本的相关度。
    使用多维度评估：
    1. 直接出现 → 0 分（已存在，无需再嵌入）
    2. 同义词/近义词 → 高相关
    3. 上下文相关词出现 → 中相关
    4. 无关联 → 低相关
    """
    text_lower = text.lower()
    kw_lower = keyword.lower()

    # 已存在
    if kw_lower in text_lower or keyword in text:
        return 0.0

    # 用 jieba 分词后检查子词匹配
    text_words = set(jieba.lcut(text))
    kw_parts = set(jieba.lcut(keyword))

    # 计算子词重叠
    overlap = len(text_words & kw_parts)
    if overlap > 0:
        return min(0.4 + overlap * 0.15, 0.85)

    # 单字匹配
    kw_chars = set(kw_lower)
    text_chars = set(text_lower)
    char_overlap = len(kw_chars & text_chars) / max(len(kw_chars), 1)
    if char_overlap > 0.5:
        return char_overlap * 0.5

    return 0.1  # 最低相关度，仍可尝试嵌入


def _match_keywords_to_section(keywords_with_weights, section_text: str,
                                top_n: int = 5) -> List[Tuple[str, float]]:
    """
    从关键词列表中选出与当前区块最相关的 top_n 个关键词。
    返回：[(keyword, relevance_score), ...]
    """
    if not section_text:
        return []

    scored = []
    for kw, weight in keywords_with_weights:
        rel = _calculate_text_relevance(section_text, kw)
        if rel > 0.05:  # 至少有点相关
            # 综合得分 = 关键词权重 × 相关度
            combined = weight * (0.5 + rel * 0.5)
            scored.append((kw, combined))

    scored.sort(key=lambda x: x[1], reverse=True)
    return scored[:top_n]


# ---------------------------------------------------------------------------
# 3.4 自然关键词嵌入策略（核心：关键词在句首/句中出现，而非末尾追加）
# ---------------------------------------------------------------------------

def _natural_weave_keyword(text: str, keyword: str, position: str = "auto") -> str:
    """
    将 keyword 自然地编织进 text 中，让关键词在显眼位置（句首/句中）出现。
    
    嵌入策略（按自然度排序）：
    1. 句首修饰：keyword + "背景下/框架下/平台上，" + text
    2. 句中替换：用 keyword 替换 text 中的泛化词
    3. 动作修饰：将 keyword 作为动词的宾语或修饰语
    4. 并列扩展：text + "与" + keyword + 相关动作
    
    position: "beginning" / "middle" / "end" / "auto"
    """
    if not text or not keyword:
        return text
    if keyword.lower() in text.lower() or keyword in text:
        return text  # 已存在

    is_cn = _is_chinese(text)

    # ====== 策略选择 ======
    if position == "auto":
        # 根据关键词类型选策略
        kw_len = len(keyword)
        if is_cn:
            if kw_len <= 4:
                position = "beginning"  # 短词 → 句首修饰最显眼
            elif kw_len <= 8:
                position = "middle"     # 中等词 → 句中融合
            else:
                position = "beginning"  # 长词 → 句首作为背景
        else:
            position = "beginning"

    # ====== 中文嵌入策略 ======
    if is_cn:
        # 策略1：句首修饰（最显眼）
        if position == "beginning":
            patterns = [
                f"基于{keyword}，{text}",
                f"在{keyword}框架下，{text}",
                f"运用{keyword}，{text}",
                f"通过{keyword}方式，{text}",
                f"以{keyword}为核心，{text}",
                f"聚焦{keyword}领域，{text}",
                f"在{keyword}方面，{text}",
                f"作为{keyword}负责人，{text}",
            ]
            # 选第一个不产生语义重复的模式
            for pattern in patterns:
                result = pattern
                # 如果 text 已经以某个动词开头，去掉模式的结尾逻辑
                if any(text.startswith(v) for v in ACTION_VERBS_CN):
                    # 保留 text 原样
                    pass
                return result

        # 策略2：句中融合
        elif position == "middle":
            # 找到第一个逗号或动词后的位置插入
            comma_pos = text.find("，")
            if comma_pos == -1:
                comma_pos = text.find(",")
            if comma_pos > 3:
                return text[:comma_pos] + f"（含{keyword}）" + text[comma_pos:]
            # 在动词之后插入
            for verb in ACTION_VERBS_CN[:10]:
                vpos = text.find(verb)
                if vpos >= 0:
                    end_pos = vpos + len(verb)
                    return text[:end_pos] + f"{keyword}相关" + text[end_pos:]
            # 无处可插，退化为句首
            return f"在{keyword}方面，{text}"

        # 策略3：末尾强化（但比旧版好——作为成就的补充）
        else:
            result = text.rstrip("。，,.、;；!！?？")
            patterns = [
                f"{result}，实现{keyword}能力突破",
                f"{result}，提升{keyword}水平",
                f"{result}，建立{keyword}体系",
                f"{result}，推动{keyword}落地",
            ]
            return patterns[0]

    # ====== 英文嵌入策略 ======
    else:
        if position == "beginning":
            patterns = [
                f"Leveraging {keyword}, {text[0].lower()}{text[1:]}" if text else f"Leveraging {keyword}",
                f"Driven by {keyword}, {text[0].lower()}{text[1:]}" if text else f"Driven by {keyword}",
                f"With expertise in {keyword}, {text[0].lower()}{text[1:]}" if text else f"Expert in {keyword}",
            ]
            return patterns[0]
        elif position == "middle":
            comma_pos = text.find(",")
            if comma_pos > 3:
                return text[:comma_pos] + f" (including {keyword})" + text[comma_pos:]
            return f"{text}, with focus on {keyword}"
        else:
            result = text.rstrip(".,;!?")
            return f"{result}, strengthening {keyword} capabilities"


def _weave_multiple_keywords(text: str, keywords: List[Tuple[str, float]],
                              max_kw: int = 3) -> str:
    """
    将多个关键词自然地编织进同一段文本。
    keywords: [(keyword, relevance_score), ...]，已按相关度排序
    
    策略：
    - 第一个（最相关）→ 句首修饰
    - 第二个 → 句中融合
    - 第三个 → 末尾强化
    - 超过3个的不硬塞
    """
    if not text or not keywords:
        return text

    result = text
    positions = ["beginning", "middle", "end"]
    used = set()

    for i, (kw, _) in enumerate(keywords[:max_kw]):
        if kw.lower() in result.lower() or kw in result:
            continue  # skip already present
        if kw in used:
            continue
        # 检查是否真的是有意义的词
        if len(kw) < 2:
            continue

        pos = positions[min(i, len(positions) - 1)]
        # 第一次用 beginning，之后用 middle
        if i > 0 and pos == "beginning":
            pos = "middle"

        result = _natural_weave_keyword(result, kw, position=pos)
        used.add(kw)

    return result


# ---------------------------------------------------------------------------
# 3.5 区块专项增强函数
# ---------------------------------------------------------------------------

def _enhance_summary_text(text: str, top_keywords: List[str],
                           action_verbs: List[str]) -> str:
    """
    增强个人总结区块：将最核心的关键词编织进总结中。
    策略：在总结中找到合适位置自然插入核心关键词。
    """
    if not text or not top_keywords:
        return text

    is_cn = _is_chinese(text)

    # 找出当前文本已覆盖的关键词
    text_lower = text.lower()
    missing = [kw for kw in top_keywords[:8]
               if kw.lower() not in text_lower and len(kw) >= 2]

    if not missing:
        return text

    result = text

    # 策略：在总结的开头或第一句之后嵌入2-3个核心关键词
    if is_cn:
        # 找到第一句话的结束位置
        first_period = result.find("。")
        first_comma = result.find("，")
        first_dot = result.find(".")

        # 优先在第一个逗号后插入
        insert_pos = first_comma if first_comma > 5 else (
            first_period if first_period > 5 else min(len(result), 30))

        if insert_pos > 5 and insert_pos < len(result):
            kw_snippet = "、".join(missing[:3])
            insert_text = f"精通{kw_snippet}"
            # 检查是否太突兀，如果是自我评价类总结，就更自然
            result = result[:insert_pos] + f"，{insert_text}" + result[insert_pos:]

    else:
        # 英文：在 after expertise/experience 之类的词后插入
        expertise_patterns = [
            (r'(?i)(expertise in|experienced in|background in|specializing in)\s+',
             f'\\1{", ".join(missing[:3])}, '),
        ]
        for pat, repl in expertise_patterns:
            if re.search(pat, result):
                result = re.sub(pat, repl, result, count=1)
                break
        else:
            # 没有找到合适位置，在开头插入
            result = f"With expertise in {', '.join(missing[:3])}, {result[0].lower()}{result[1:]}"

    return result


def _enhance_skill_text(text: str, hard_skills: List[Tuple[str, float]],
                         max_add: int = 5) -> str:
    """
    增强技能列表区块：直接补充缺失的硬技能关键词。
    技能列表通常是逗号/顿号分隔的，直接追加即可。
    """
    if not text or not hard_skills:
        return text

    is_cn = _is_chinese(text)
    text_lower = text.lower()

    # 找出缺失的硬技能
    missing = []
    for kw, weight in hard_skills:
        if kw.lower() not in text_lower and len(kw) >= 2:
            missing.append(kw)
            if len(missing) >= max_add:
                break

    if not missing:
        return text

    # 追加到末尾
    result = text.rstrip("。，,.、;；!！?？/ ")
    sep = "、" if is_cn else ", "
    result += sep + sep.join(missing)

    return result


def _enhance_experience_text(text: str,
                              relevant_kw: List[Tuple[str, float]],
                              top_keywords: List[str]) -> str:
    """
    增强工作经历/项目经验单条文本：
    1. 动词强化（弱动词→强动词）
    2. 匹配最相关的JD关键词，自然编织进去
    3. 确保一条经历至少有1-2个JD核心关键词可见
    """
    if not text or not text.strip():
        return text

    is_cn = _is_chinese(text)
    result = text

    # ---- 步骤1：动词强化 ----
    if is_cn:
        verb_upgrades = [
            ("做了", "主导完成"), ("做了很多", "主导多项"),
            ("参与", "深度参与"), ("帮助", "支撑"),
            ("做", "完成"), ("搞", "推进"), ("弄", "落地实施"),
            ("处理", "高效处理"), ("写", "撰写"),
            ("管理", "统筹管理"), ("了解", "熟练掌握"),
            ("会使用", "精通"), ("用过", "熟练运用"),
            ("知道", "深入理解"),
        ]
        for weak, strong in verb_upgrades:
            if weak in result:
                result = result.replace(weak, strong, 1)
                break
    else:
        verb_upgrades_en = [
            ("did", "executed"), ("made", "developed"),
            ("helped", "supported"), ("worked on", "led"),
            ("was responsible for", "spearheaded"),
            ("participated in", "drove"), ("handled", "managed"),
            ("used", "leveraged"), ("wrote", "authored"),
        ]
        for weak, strong in verb_upgrades_en:
            pattern = re.compile(weak, re.IGNORECASE)
            if pattern.search(result):
                result = pattern.sub(strong, result, count=1)
                break

    # ---- 步骤2：智能关键词编织 ----
    # 筛选相关且不在文本中的关键词
    text_lower = result.lower()
    to_weave = []
    # 先从匹配的关键词中选
    if relevant_kw:
        for kw, score in relevant_kw:
            if kw.lower() not in text_lower and len(kw) >= 2:
                to_weave.append((kw, score))
            if len(to_weave) >= 3:
                break
    # 不够的话从 top_keywords 补充
    if len(to_weave) < 2:
        for kw in top_keywords:
            if kw.lower() not in text_lower and len(kw) >= 2:
                # 检查是否已经被 to_weave 覆盖
                if kw not in [t[0] for t in to_weave]:
                    to_weave.append((kw, 0.3))
            if len(to_weave) >= 3:
                break

    if to_weave:
        result = _weave_multiple_keywords(result, to_weave, max_kw=2)

    return result


# ---------------------------------------------------------------------------
# 3.6 改写主函数（LLM 优先，jieba 兜底）
# ---------------------------------------------------------------------------

def rewrite_resume(matched_doc_path: str, job_description: str, job_analysis: Dict) -> str:
    """
    智能简历改写引擎 — 双引擎架构：
    1. 如果配置了 LLM API Key → 使用大模型推理改写（自然、无重复）
    2. 否则 → 使用本地 jieba + 规则引擎兜底
    
    返回：优化后简历的文件路径。
    """
    # ====== 检测 LLM 是否可用 ======
    if _llm_available():
        return _rewrite_resume_with_llm(matched_doc_path, job_description, job_analysis)

    # ====== 本地规则引擎（jieba，作为兜底方案） ======
    st.info("💡 未配置 LLM API Key，使用本地规则引擎改写（侧边栏可配置 LLM 获得更好效果）")
    # ====== 1. 深度分析 JD ======
    with st.spinner("🔬 正在深度分析岗位核心要求..."):
        jd_deep = _analyze_jd_deep(job_description, job_analysis)
        top_keywords = jd_deep["top_keywords"]
        hard_skills = jd_deep["hard_skills"]
        soft_skills = jd_deep["soft_skills"]
        all_weighted = jd_deep["all_weighted"]

    # ====== 2. 读取文档并解析结构 ======
    with st.spinner("📖 正在读取并分析简历结构..."):
        doc = Document(matched_doc_path)

    # ====== 3. 收集所有段落并识别区块 ======
    with st.spinner("🔄 正在智能匹配并改写简历..."):
        total_modified = 0
        current_section = "header"  # header / summary / skills / experience / education / other

        # ---------- 第一遍：识别区块 ----------
        section_boundaries = []  # [(para_index, section_type), ...]
        para_index = 0
        para_texts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            para_texts.append(text)
            section_type = _detect_section_type(text)
            if section_type != "content":
                section_boundaries.append((para_index, section_type))
            para_index += 1

        # 建立段落到区块的映射
        para_section_map = {}
        if section_boundaries:
            for i in range(len(section_boundaries)):
                start_idx = section_boundaries[i][0]
                end_idx = section_boundaries[i+1][0] if i+1 < len(section_boundaries) else len(para_texts)
                for j in range(start_idx, min(end_idx, len(para_texts))):
                    para_section_map[j] = section_boundaries[i][1]

        # ---------- 第二遍：逐段落优化 ----------
        # 收集完整的工作经历区块文本（用于关键词匹配）
        experience_full_text = ""
        summary_full_text = ""
        for i, text in enumerate(para_texts):
            sec = para_section_map.get(i, "other")
            if sec == "experience":
                experience_full_text += text + " "
            elif sec == "summary":
                summary_full_text += text + " "

        para_idx = 0
        for para in doc.paragraphs:
            original_text = para.text.strip()
            if not original_text:
                para_idx += 1
                continue

            section = para_section_map.get(para_idx, "other")
            is_cn = _is_chinese(original_text)

            # 跳过标题行本身
            if _is_section_header(original_text):
                para_idx += 1
                continue

            new_text = original_text

            # ---- 区块专属策略 ----
            if section == "summary":
                # 个人总结：嵌入最核心的关键词
                if len(original_text) > 10:
                    new_text = _enhance_summary_text(
                        original_text, top_keywords, jd_deep["action_verbs"])

            elif section == "skills":
                # 技能列表：补充缺失的硬技能
                new_text = _enhance_skill_text(original_text, hard_skills, max_add=5)

            elif section == "experience":
                # 工作经历/项目经验：动词强化 + 智能关键词编织
                # 判断是否是内容行（非空行、非纯日期/公司名）
                is_content = len(original_text) > 8
                is_action = any(
                    verb in original_text[:15]
                    for verb in ACTION_VERBS_CN + ["led", "managed", "developed",
                                                    "designed", "built", "created"]
                )
                # 也检测以项目符号开头的行
                starts_with_bullet = original_text[0] in "•●·-–—>*○◆▪▸"

                if is_content and (is_action or starts_with_bullet or len(original_text) > 15):
                    # 为这段经历匹配最相关的 JD 关键词
                    relevant = _match_keywords_to_section(
                        [(kw, all_weighted.get(kw, 0.5)) for kw in top_keywords],
                        original_text, top_n=5)
                    new_text = _enhance_experience_text(
                        original_text, relevant, top_keywords)

            elif section == "other":
                # 未识别区块：如果包含动作动词则按经历处理
                if len(original_text) > 10:
                    has_action = any(
                        verb in original_text[:15]
                        for verb in ACTION_VERBS_CN + ["led", "managed", "developed",
                                                        "designed", "built", "created"])
                    starts_with_bullet = original_text[0] in "•●·-–—>*○◆▪▸"
                    if has_action or starts_with_bullet:
                        relevant = _match_keywords_to_section(
                            [(kw, all_weighted.get(kw, 0.5)) for kw in top_keywords],
                            original_text, top_n=5)
                        new_text = _enhance_experience_text(
                            original_text, relevant, top_keywords)

            # ---- 写回段落 ----
            if new_text != original_text:
                if para.runs:
                    first_run = para.runs[0]
                    for i in range(1, len(para.runs)):
                        para.runs[i].text = ""
                    first_run.text = new_text
                else:
                    para.add_run(new_text)
                total_modified += 1

            para_idx += 1

        # ---------- 优化表格（表格中的内容，通常也是经历/技能） ----------
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        original_text = para.text.strip()
                        if not original_text or len(original_text) < 8:
                            continue

                        is_cn = _is_chinese(original_text)
                        is_action = any(
                            verb in original_text[:15]
                            for verb in ACTION_VERBS_CN + ["led", "managed",
                                                            "developed", "designed"])
                        starts_with_bullet = original_text[0] in "•●·-–—>*○◆▪▸"

                        if is_action or starts_with_bullet or len(original_text) > 20:
                            relevant = _match_keywords_to_section(
                                [(kw, all_weighted.get(kw, 0.5)) for kw in top_keywords],
                                original_text, top_n=5)
                            new_text = _enhance_experience_text(
                                original_text, relevant, top_keywords)

                            if new_text != original_text:
                                if para.runs:
                                    first_run = para.runs[0]
                                    for i in range(1, len(para.runs)):
                                        para.runs[i].text = ""
                                    first_run.text = new_text
                                else:
                                    para.add_run(new_text)
                                total_modified += 1

    # ====== 4. 保存新文件 ======
    with st.spinner("💾 正在保存优化后的简历..."):
        output_dir = os.path.dirname(matched_doc_path)
        original_name = os.path.basename(matched_doc_path)
        name_without_ext, ext = os.path.splitext(original_name)
        new_name = f"优化后_{name_without_ext}{ext}"
        new_path = os.path.join(output_dir, new_name)

        counter = 1
        while os.path.exists(new_path):
            new_name = f"优化后_{name_without_ext}_v{counter}{ext}"
            new_path = os.path.join(output_dir, new_name)
            counter += 1

        doc.save(new_path)

    # 轻量结果统计
    st.caption(f"📊 改写统计：共优化 {total_modified} 处，核心关键词 {len(top_keywords)} 个已智能嵌入")

    return new_path


# ============================================================================
# Streamlit UI 布局
# ============================================================================

def render_sidebar():
    """渲染侧边栏"""
    st.sidebar.markdown(
        '<h1 style="font-size:22px; font-weight:600;">智能简历工具</h1>'
        '<p style="font-size:13px; color:#86868b; margin-top:-0.8rem;">本地版 · 数据不出设备</p>',
        unsafe_allow_html=True,
    )
    st.sidebar.markdown("---")

    # ---------- 步骤导航（苹果风格圆点） ----------
    st.sidebar.markdown("### 操作步骤")
    steps = st.session_state.get("step_completed", {})
    step_labels = [
        ("step1", "输入岗位要求"),
        ("step2", "分析岗位核心要求"),
        ("step3", "选取简历文件夹"),
        ("step4", "搜索最佳匹配简历"),
        ("step5", "本地改写简历"),
        ("step6", "一键转 PDF"),
    ]
    for sid, slabel in step_labels:
        done = steps.get(sid)
        dot_class = "step-dot-done" if done else "step-dot-todo"
        icon = "✓" if done else str(step_labels.index((sid, slabel)) + 1)
        text_class = "sidebar-step-done" if done else ""
        st.sidebar.markdown(
            f'<div class="sidebar-step {text_class}">'
            f'<span class="step-dot {dot_class}">{icon}</span>'
            f'<span>{slabel}</span>'
            f'</div>',
            unsafe_allow_html=True,
        )

    st.sidebar.markdown("---")
    st.sidebar.caption("纯本地处理 · 无需联网 · 无需 API Key")

    # ---------- LLM 配置（可选，用于智能改写） ----------
    st.sidebar.markdown("---")
    st.sidebar.markdown("### 🤖 大模型改写（可选）")
    st.sidebar.caption("配置后可获得更自然、无重复的改写效果。留空则使用本地规则引擎。")

    with st.sidebar.expander("⚙️ 配置 LLM API", expanded=False):
        # 快速预设
        preset = st.selectbox(
            "服务商预设",
            ["手动配置", "DeepSeek（推荐·便宜好用）", "OpenAI", "Moonshot（Kimi）", "零一万物", "智谱（GLM）", "Ollama 本地"],
            key="llm_preset",
        )

        preset_configs = {
            "DeepSeek（推荐·便宜好用）": {
                "base_url": "https://api.deepseek.com",
                "model": "deepseek-chat",
            },
            "OpenAI": {
                "base_url": "https://api.openai.com/v1",
                "model": "gpt-4o",
            },
            "Moonshot（Kimi）": {
                "base_url": "https://api.moonshot.cn/v1",
                "model": "moonshot-v1-8k",
            },
            "零一万物": {
                "base_url": "https://api.lingyiwanwu.com/v1",
                "model": "yi-large",
            },
            "智谱（GLM）": {
                "base_url": "https://open.bigmodel.cn/api/paas/v4",
                "model": "glm-4-flash",
            },
            "Ollama 本地": {
                "base_url": "http://localhost:11434/v1",
                "model": "qwen2.5:7b",
            },
        }

        if preset != "手动配置" and preset in preset_configs:
            cfg = preset_configs[preset]
            st.session_state["llm_base_url"] = cfg["base_url"]
            st.session_state["llm_model"] = cfg["model"]

        llm_api_key = st.text_input(
            "API Key",
            value=st.session_state.get("llm_api_key", ""),
            type="password",
            placeholder="留空则不启用大模型改写",
        )
        if llm_api_key:
            st.session_state["llm_api_key"] = llm_api_key

        llm_base_url = st.text_input(
            "Base URL",
            value=st.session_state.get("llm_base_url", "https://api.deepseek.com"),
            placeholder="https://api.deepseek.com",
        )
        if llm_base_url:
            st.session_state["llm_base_url"] = llm_base_url

        llm_model = st.text_input(
            "Model 名称",
            value=st.session_state.get("llm_model", "deepseek-chat"),
            placeholder="deepseek-chat",
        )
        if llm_model:
            st.session_state["llm_model"] = llm_model

        if llm_api_key:
            st.success(f"✅ LLM 已配置：{st.session_state['llm_model']}")
            st.caption("🔒 Key 仅保存在你本次浏览器会话中，关闭页面或刷新即清除，不会写入服务器磁盘。")
        else:
            st.info("💡 未配置 API Key，将使用本地规则引擎改写")


# ============================================================================
# Word → PDF 转换引擎
# ============================================================================

def _find_libreoffice() -> Optional[str]:
    """查找 LibreOffice 可执行文件路径"""
    # macOS / Linux（云端部署通过 packages.txt 用 apt 安装 libreoffice）
    known_paths = [
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        "/usr/local/bin/soffice",
        "/usr/bin/soffice",
        "/usr/lib/libreoffice/program/soffice",
    ]
    for p in known_paths:
        if os.path.exists(p):
            return p
    # 检查 PATH
    soffice = shutil.which("soffice")
    if soffice:
        return soffice
    return None


def _find_cjk_font() -> Tuple[str, str, str, str]:
    """查找系统 CJK 字体，返回 (regular_path, bold_path, regular_name, bold_name)"""
    candidates = [
        # PingFang (macOS 首选)
        ("/System/Library/Fonts/PingFang.ttc", "PingFang-SC-Regular", "PingFang-SC-Bold"),
        # STHeiti
        ("/System/Library/Fonts/STHeiti Light.ttc", "STHeiti-Regular", "STHeiti-Bold"),
        ("/System/Library/Fonts/STHeiti Medium.ttc", "STHeiti-Regular", "STHeiti-Bold"),
        # Hiragino
        ("/System/Library/Fonts/Hiragino Sans GB.ttc", "HiraginoSansGB-W3", "HiraginoSansGB-W6"),
        # 回退字体
        ("/Library/Fonts/Arial Unicode.ttf", "ArialUnicodeMS", "ArialUnicodeMS"),
        # Linux 云端（packages.txt 用 apt 安装 fonts-noto-cjk）
        ("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc", "NotoSansCJK-Regular", "NotoSansCJK-Bold"),
        ("/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc", "NotoSansCJK-Regular", "NotoSansCJK-Bold"),
    ]
    for fp, reg, bold in candidates:
        if os.path.exists(fp):
            return fp, fp, reg, bold
    # 扫描系统字体目录
    for base in ["/System/Library/Fonts", "/Library/Fonts",
                 "/usr/share/fonts/opentype/noto", "/usr/share/fonts/truetype/noto",
                 "/usr/share/fonts"]:
        if os.path.isdir(base):
            for fname in sorted(os.listdir(base)):
                if fname.endswith((".ttc", ".ttf", ".otf")):
                    fp = os.path.join(base, fname)
                    return fp, fp, "CJK-Regular", "CJK-Bold"
    raise RuntimeError("找不到 CJK 字体")


def _pdf_fallback_reportlab(docx_path: str, pdf_path: str) -> None:
    """用 reportlab 生成 PDF，保留字体/大小/粗体/对齐/颜色/表格等格式"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak,
    )
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    font_path, bold_path, reg_name, bold_name = _find_cjk_font()

    # 注册 CJK 字体
    pdfmetrics.registerFont(TTFont("CJK-Regular", font_path, subfontIndex=0))
    # 尝试注册粗体（同一 ttc 的不同 subfont）
    pdfmetrics.registerFont(TTFont("CJK-Bold", bold_path, subfontIndex=0))

    doc = Document(docx_path)

    # 构建样式表
    def build_style(para, base_size=10):
        """根据段落的 run 属性构建 ParagraphStyle"""
        font_name = "CJK-Bold" if any(r.bold for r in para.runs if r.bold) else "CJK-Regular"
        font_size = base_size

        # 取第一个 run 的属性作为段落级样式
        for run in para.runs:
            if run.font.size:
                font_size = run.font.size.pt
            if run.bold:
                font_name = "CJK-Bold"
            break

        # 对齐
        align_map = {
            0: TA_LEFT,
            1: TA_CENTER,
            2: TA_RIGHT,
            3: TA_JUSTIFY,
        }
        alignment = align_map.get(para.alignment, TA_LEFT)

        # 颜色
        text_color = None
        for run in para.runs:
            if run.font.color and run.font.color.rgb:
                text_color = HexColor(f"#{run.font.color.rgb}")
                break

        style = ParagraphStyle(
            f"style_{id(para)}",
            fontName=font_name,
            fontSize=font_size,
            leading=font_size * 1.5,
            alignment=alignment,
            spaceAfter=2 * mm,
            spaceBefore=0,
        )
        if text_color:
            style.textColor = text_color
        return style

    # 识别标题段落（加粗 + 字号较大）
    def is_heading(para):
        for run in para.runs:
            if run.bold and run.font.size and run.font.size.pt >= 12:
                return True
        return False

    story = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            story.append(Spacer(1, 4 * mm))
            continue

        style = build_style(para)
        # 标题样式增强
        if is_heading(para):
            style.fontSize = max(style.fontSize, 14)
            style.leading = style.fontSize * 1.8
            style.spaceBefore = 6 * mm
            style.spaceAfter = 4 * mm

        # 段落中可能有多个 run（不同格式），需要拼成富文本
        parts = []
        for run in para.runs:
            if not run.text:
                continue
            run_bold = run.bold or (style.fontName == "CJK-Bold")
            run_size = run.font.size.pt if run.font.size else style.fontSize
            fn = "CJK-Bold" if run_bold else "CJK-Regular"
            # 转义 XML 特殊字符
            escaped = run.text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            parts.append(f'<font face="{fn}" size="{run_size}">{escaped}</font>')

        if parts:
            html_text = "".join(parts)
        else:
            html_text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

        try:
            story.append(Paragraph(html_text, style))
        except Exception:
            story.append(Paragraph(text.replace("&", "&amp;").replace("<", "&lt;"), style))

    # 处理表格
    for table in doc.tables:
        story.append(Spacer(1, 4 * mm))
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = "\n".join(p.text for p in cell.paragraphs if p.text.strip())
                row_data.append(cell_text)
            table_data.append(row_data)

        if table_data:
            # 计算列宽
            num_cols = max(len(r) for r in table_data)
            col_width = (A4[0] - 30 * mm) / num_cols
            t = Table(table_data, colWidths=[col_width] * num_cols)
            t.setStyle(TableStyle([
                ("FONTNAME", (0, 0), (-1, -1), "CJK-Regular"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.5, (180, 180, 180)),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]))
            story.append(t)
            story.append(Spacer(1, 4 * mm))

    # 构建 PDF
    pdf_doc = SimpleDocTemplate(
        pdf_path,
        pagesize=A4,
        leftMargin=15 * mm,
        rightMargin=15 * mm,
        topMargin=15 * mm,
        bottomMargin=15 * mm,
    )
    pdf_doc.build(story)


def convert_word_to_pdf(docx_path: str, output_dir: Optional[str] = None) -> Tuple[str, str]:
    """将 Word (.docx) 文件转换为 PDF。

    返回: (pdf_path, method_used)
    method_used 可能值: 'word' | 'libreoffice' | 'reportlab'
    """
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"文件不存在: {docx_path}")

    if output_dir is None:
        output_dir = os.path.dirname(docx_path)

    base_name = os.path.splitext(os.path.basename(docx_path))[0]
    pdf_path = os.path.join(output_dir, f"{base_name}.pdf")

    # 方案 1: LibreOffice headless（无需授权，静默运行）
    lo_path = _find_libreoffice()
    if lo_path:
        try:
            subprocess.run(
                [lo_path, "--headless", "--convert-to", "pdf",
                 "--outdir", output_dir, docx_path],
                capture_output=True, text=True, timeout=60,
            )
            expected = os.path.join(output_dir, f"{base_name}.pdf")
            if os.path.exists(expected):
                if expected != pdf_path:
                    if os.path.exists(pdf_path):
                        os.remove(pdf_path)
                    os.rename(expected, pdf_path)
                return pdf_path, "libreoffice"
        except Exception:
            pass

    # 方案 2: Microsoft Word (macOS，格式完美，但需授权)
    if sys.platform == "darwin" and os.path.exists("/Applications/Microsoft Word.app"):
        try:
            from docx2pdf import convert as docx2pdf_convert
            import threading
            result = {"ok": False, "err": None}

            def _convert():
                try:
                    docx2pdf_convert(docx_path, pdf_path)
                    result["ok"] = True
                except Exception as e:
                    result["err"] = e

            t = threading.Thread(target=_convert, daemon=True)
            t.start()
            t.join(timeout=30)  # 30 秒超时（Word 弹授权框会卡住）

            if result["ok"] and os.path.exists(pdf_path):
                return pdf_path, "word"
        except Exception:
            pass

    # 方案 3: reportlab 纯 Python（保留字体/大小/粗体/表格/对齐/颜色）
    _pdf_fallback_reportlab(docx_path, pdf_path)
    return pdf_path, "reportlab"


def render_main_area():
    """渲染主区域"""
    # 苹果风格 Hero 标题
    st.markdown(
        '<h1 style="text-align:center; font-size:56px; margin-top:1rem;">智能简历匹配与改写</h1>',
        unsafe_allow_html=True,
    )
    st.markdown(
        '<p style="text-align:center; font-size:21px; color:#86868b; font-weight:400; margin-top:-0.5rem; margin-bottom:0.5rem;">'
        '上传岗位描述，分析核心要求，匹配最佳简历，智能改写优化。'
        '</p>',
        unsafe_allow_html=True,
    )
    st.markdown(
        '<p style="text-align:center; font-size:15px; color:#86868b;">'
        '🔒 所有数据处理在本地完成 · 无需联网 · 无需下载模型 · 秒级响应'
        '</p>',
        unsafe_allow_html=True,
    )
    st.markdown("---")

    # ========================================================================
    # 步骤1：输入岗位要求
    # ========================================================================
    st.markdown("## 📝 步骤 1：输入岗位要求")
    st.markdown("直接粘贴岗位描述文本，或上传岗位截图/文件（支持 PNG/JPG/PDF/DOCX）")

    col_text, col_upload = st.columns([3, 2])

    with col_text:
        job_text_input = st.text_area(
            "岗位描述文本",
            value=st.session_state.get("job_description_text", ""),
            height=250,
            placeholder="在此粘贴岗位描述(Job Description)...\n\n例如：\n职位名称：ERP实施项目经理\n职责：\n1. 负责ERP系统的实施与项目管理\n2. 协调业务部门与技术团队\n要求：\n- 5年以上ERP实施经验\n- 熟悉SAP/Oracle等主流ERP系统",
            key="job_text_area",
        )
        if job_text_input:
            st.session_state["job_description_text"] = job_text_input

    with col_upload:
        st.markdown("#### 或者上传文件")
        uploaded_file = st.file_uploader(
            "上传岗位描述文件",
            type=["png", "jpg", "jpeg", "pdf", "docx"],
            help="支持：截图(PNG/JPG)、PDF、Word(.docx)",
        )

        if uploaded_file is not None:
            try:
                extracted_text = extract_text_from_upload(uploaded_file)
                if extracted_text and not extracted_text.startswith("❌") and not extracted_text.startswith("⚠️"):
                    st.session_state["job_description_text"] = extracted_text
                    st.success(f"✅ 成功提取 {len(extracted_text)} 个字符")
                    with st.expander("👀 预览提取的文字"):
                        st.text(extracted_text[:1500])
                    st.rerun()
                elif extracted_text.startswith("⚠️"):
                    st.warning(extracted_text)
                else:
                    st.error(extracted_text)
            except Exception as e:
                st.error(f"❌ 文件处理失败: {e}")

    if st.session_state.get("job_description_text", "").strip():
        st.session_state["step_completed"]["step1"] = True
        st.success(f"✅ 步骤1完成 — 岗位描述已输入 ({len(st.session_state['job_description_text'])} 字符)")
    else:
        st.session_state["step_completed"]["step1"] = False

    st.markdown("---")

    # ========================================================================
    # 步骤2：本地分析岗位核心要求
    # ========================================================================
    st.markdown("## 🔍 步骤 2：分析岗位核心要求")
    st.caption("使用 jieba 分词 + TF-IDF + 技能词典 + 关键词分类 纯本地分析")

    col_btn2, col_status2 = st.columns([1, 4])
    with col_btn2:
        analyze_btn = st.button(
            "🚀 分析岗位要求",
            type="primary",
            disabled=not st.session_state["step_completed"]["step1"],
            use_container_width=True,
        )

    if analyze_btn:
        try:
            with st.spinner("正在使用本地 NLP 分析岗位要求..."):
                result = get_job_keywords(st.session_state["job_description_text"])
                st.session_state["job_analysis_result"] = result
                st.session_state["step_completed"]["step2"] = True
            st.rerun()
        except Exception as e:
            st.error(f"❌ 分析失败: {e}")

    # 显示分析结果
    if st.session_state.get("job_analysis_result"):
        analysis = st.session_state["job_analysis_result"]

        st.markdown("### 🎯 分析结果")

        # 核心技能
        st.markdown("#### 💪 核心技能")
        if analysis.get("skills"):
            cols = st.columns(min(len(analysis["skills"]), 4))
            for i, skill in enumerate(analysis["skills"]):
                with cols[i % 4]:
                    st.markdown(
                        f'<div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); '
                        f'padding: 8px 14px; border-radius: 20px; color: white; '
                        f'margin: 4px 0; font-size: 14px; text-align: center;">{skill}</div>',
                        unsafe_allow_html=True,
                    )
        else:
            st.info("暂无技能数据")

        st.markdown("---")

        # 核心要求/职责
        st.markdown("#### 📋 核心要求/职责")
        if analysis.get("requirements"):
            for i, req in enumerate(analysis["requirements"], 1):
                st.markdown(
                    f'<div style="background: #f0f8ff; border-left: 4px solid #2196F3; '
                    f'padding: 10px 15px; margin: 8px 0; border-radius: 4px;">'
                    f'<strong>{i}.</strong> {req}</div>',
                    unsafe_allow_html=True,
                )
        else:
            st.info("暂无要求数据")

        st.markdown("---")

        # 重要关键词
        st.markdown("#### 🏷️ 重要关键词")
        if analysis.get("keywords"):
            tags_html = ""
            colors = ["#e3f2fd", "#fce4ec", "#e8f5e9", "#fff3e0", "#f3e5f5", "#e0f7fa"]
            for i, kw in enumerate(analysis["keywords"]):
                color = colors[i % len(colors)]
                tags_html += (
                    f'<span style="display: inline-block; background: {color}; '
                    f'padding: 4px 12px; margin: 3px; border-radius: 12px; '
                    f'font-size: 13px; border: 1px solid #ddd;">{kw}</span>'
                )
            st.markdown(
                f'<div style="line-height: 2.2;">{tags_html}</div>',
                unsafe_allow_html=True,
            )
        else:
            st.info("暂无关键词数据")

        st.markdown("---")

    # ========================================================================
    # 步骤3：上传简历文件
    # ========================================================================
    st.markdown("## 📂 步骤 3：上传简历文件")
    st.caption("上传一份或多份 .docx 简历，工具会自动扫描并匹配（文件只保存在你本次会话中）")

    uploaded_resumes = st.file_uploader(
        "上传简历文件（.docx，可多选）",
        type=["docx"],
        accept_multiple_files=True,
        disabled=not st.session_state["step_completed"].get("step2"),
        key="resume_uploader",
        label_visibility="collapsed",
    )

    if uploaded_resumes:
        resume_dir = st.session_state.get("resume_folder", "")
        if not resume_dir or not os.path.isdir(resume_dir):
            resume_dir = tempfile.mkdtemp(prefix="resumes_")
            st.session_state["resume_folder"] = resume_dir
        for f in uploaded_resumes:
            safe_name = os.path.basename(f.name)
            if not safe_name.lower().endswith(".docx"):
                continue
            with open(os.path.join(resume_dir, safe_name), "wb") as out:
                out.write(f.getbuffer())

    # 显示已上传文件状态
    current_folder = st.session_state.get("resume_folder", "")
    if current_folder and os.path.isdir(current_folder):
        docx_files = [
            f for f in os.listdir(current_folder)
            if f.lower().endswith(".docx") and not f.startswith("~$")
        ]
        if docx_files:
            st.success(f"✅ 已上传 {len(docx_files)} 份简历 (.docx)")
            if len(docx_files) <= 12:
                cols = st.columns(min(len(docx_files), 3))
                for i, f in enumerate(docx_files):
                    with cols[i % 3]:
                        st.caption(f"📄 {f}")
            else:
                with st.expander(f"👀 查看全部 {len(docx_files)} 份简历文件名"):
                    for f in docx_files:
                        st.caption(f"📄 {f}")
            st.session_state["step_completed"]["step3"] = True
        else:
            st.session_state["step_completed"]["step3"] = False
    else:
        st.session_state["step_completed"]["step3"] = False
        if st.session_state["step_completed"].get("step2"):
            st.info("💡 请上传存放 .docx 简历的文件（可一次选择多个文件）")

    st.markdown("---")

    # ========================================================================
    # 步骤4：搜索最佳匹配简历
    # ========================================================================
    st.markdown("## 🎯 步骤 4：搜索最佳匹配简历")
    st.caption("纯本地 jieba 关键词匹配 — 秒级响应，无需下载模型")

    col_btn4, _ = st.columns([1, 4])
    with col_btn4:
        match_disabled = not st.session_state["step_completed"].get("step3", False)
        match_btn = st.button(
            "🔍 搜索最佳匹配简历",
            type="primary",
            disabled=match_disabled,
            use_container_width=True,
        )

    if match_btn:
        folder = st.session_state.get("resume_folder", "")
        if not os.path.isdir(folder):
            st.error(f"❌ 简历文件夹无效: {folder}")
        else:
            try:
                results = find_best_resume(st.session_state["job_description_text"], folder, top_n=3)
                st.session_state["matched_resumes_list"] = results
                st.session_state["selected_resume_index"] = 0
                st.session_state["step_completed"]["step4"] = True
                st.rerun()
            except Exception as e:
                st.error(f"❌ 匹配失败: {e}")

    # 显示 Top 3 匹配结果
    matched_list = st.session_state.get("matched_resumes_list", [])
    if matched_list:
        st.markdown("### 🏆 匹配结果（Top 3）")
        st.caption("选择一份简历，然后进入步骤 5 改写")

        # 构建选项标签
        options = []
        for i, r in enumerate(matched_list):
            pct = r["similarity"] * 100
            medal = ["🥇", "🥈", "🥉"][i] if i < 3 else "📄"
            options.append(f"{medal} {r['name']} — 匹配度 {pct:.1f}%")

        selected_idx = st.radio(
            "请选择要改写的简历",
            options=range(len(options)),
            format_func=lambda i: options[i],
            index=st.session_state.get("selected_resume_index", 0),
            key="resume_selector",
        )
        st.session_state["selected_resume_index"] = selected_idx

        # 三列卡片展示
        cols = st.columns(3)
        for i, r in enumerate(matched_list):
            pct = r["similarity"] * 100
            medal = ["🥇", "🥈", "🥉"][i]

            if pct >= 70:
                bar_color = "#4CAF50"
                bg = "#e8f5e9"
            elif pct >= 50:
                bar_color = "#FF9800"
                bg = "#fff3e0"
            else:
                bar_color = "#F44336"
                bg = "#ffebee"

            is_selected = (i == selected_idx)
            border = "3px solid #2196F3" if is_selected else "1px solid #ddd"

            with cols[i]:
                st.markdown(
                    f'<div style="background:{bg}; padding:16px; border-radius:12px; '
                    f'border:{border}; height:100%;">'
                    f'<h4 style="margin:0 0 4px 0;">{medal} 第{i+1}名</h4>'
                    f'<p style="margin:0 0 8px 0; font-size:13px; color:#666; '
                    f'word-break:break-all;">{r["name"]}</p>'
                    f'<p style="margin:0 0 8px 0; font-size:20px; font-weight:bold; '
                    f'color:{bar_color};">{pct:.1f}%</p>'
                    f'<div style="background:#e0e0e0; border-radius:6px; height:8px; '
                    f'width:100%; margin-bottom:10px;">'
                    f'<div style="background:{bar_color}; border-radius:6px; height:8px; '
                    f'width:{min(pct,100)}%;"></div></div>'
                    f'</div>',
                    unsafe_allow_html=True,
                )

        # 预览选中简历
        selected = matched_list[selected_idx]
        with st.expander(f"👀 预览选中简历：{selected['name']}（前1000字）"):
            st.text(selected["preview"])

        st.markdown("---")

    # ========================================================================
    # 步骤5：智能改写简历
    # ========================================================================
    st.markdown("## ✨ 步骤 5：智能改写简历")
    if _llm_available():
        st.caption(f"🤖 大模型引擎：{st.session_state['llm_model']} — 自然表达，杜绝模板重复")
    else:
        st.caption("💡 本地规则引擎 — 侧边栏配置 LLM API Key 可获得更自然的改写效果")

    col_btn5, _ = st.columns([1, 4])
    with col_btn5:
        rewrite_disabled = not st.session_state["step_completed"].get("step4", False)
        btn_label = "🤖 LLM 智能改写" if _llm_available() else "🪄 本地改写简历"
        rewrite_btn = st.button(
            btn_label,
            type="primary",
            disabled=rewrite_disabled,
            use_container_width=True,
        )

    if rewrite_btn:
        # 使用用户选择的简历
        matched_list = st.session_state.get("matched_resumes_list", [])
        idx = st.session_state.get("selected_resume_index", 0)
        if idx < len(matched_list):
            selected = matched_list[idx]
            try:
                new_path = rewrite_resume(
                    selected["path"],
                    st.session_state["job_description_text"],
                    st.session_state.get("job_analysis_result", {}),
                )
                st.session_state["optimized_resume_path"] = new_path
                st.session_state["optimized_resume_name"] = os.path.basename(new_path)
                st.session_state["step_completed"]["step5"] = True
                st.rerun()
            except Exception as e:
                st.error(f"❌ 简历改写失败: {e}")
        else:
            st.error("❌ 请先选择一份简历")

    # 显示下载按钮 + 改写说明
    if st.session_state.get("optimized_resume_path") and os.path.exists(
        st.session_state["optimized_resume_path"]
    ):
        st.success("🎉 改写完成!")

        with open(st.session_state["optimized_resume_path"], "rb") as f:
            file_data = f.read()

        st.download_button(
            label=f"📥 下载改写后简历 ({st.session_state['optimized_resume_name']})",
            data=file_data,
            file_name=st.session_state["optimized_resume_name"],
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

        if _llm_available():
            st.info(f"""
            **🤖 LLM 智能改写引擎**
            
            🧠 **模型**：`{st.session_state['llm_model']}`
            
            🎯 **改写策略：**
            - 完整理解简历和 JD 的上下文语义
            - 关键词自然融入对应区块（句首/句中显眼位置）
            - 每段经历使用不同句式，杜绝重复
            - 弱动词自动升级为强动词
            - 不编造经历，只优化表达
            
            📌 生成的文件名包含 `_LLM` 后缀以区分来源
            
            - 保留原始 Word 格式（字体、颜色、表格、排版）
            - 教育背景、联系方式等保持原样
            """)
        else:
            st.info("""
            **本地规则引擎说明：**
            
            🔬 **深度JD分析**：关键词分类为硬技能/软技能/产品工具/行业术语，各赋权重
            
            🧩 **简历区块识别**：自动识别个人总结、技能列表、工作经历、教育背景
            
            🎯 **智能关键词自然嵌入**（非末尾追加）：
            - 句首修饰："基于SAP平台，负责ERP系统实施" 
            - 句中融合："统筹管理项目交付（含敏捷开发）"
            - 技能补充：直接在技能清单中补充缺失的硬技能
            
            📌 **核心原则**：关键词出现在显眼位置，一眼扫过即见核心匹配点
            
            💡 **提示**：侧边栏配置 LLM API Key 可获得更自然、无重复的改写效果
            
            - 保留原始 Word 格式（字体、颜色、表格、排版）
            - 教育背景、联系方式等保持原样不修改
            """)

        st.markdown(
            f"> ✅ 改写后简历已保存至：`{st.session_state['optimized_resume_path']}`"
        )

    st.markdown("---")

    # ========================================================================
    # 步骤6：一键转 PDF
    # ========================================================================
    st.markdown("## 📑 步骤 6：一键转 PDF")
    st.caption("将改写后的 Word 简历一键转换为 PDF，或上传其他 Word 文件转 PDF")

    # ---- 选项 A：转换步骤 5 生成的简历 ----
    has_optimized = (
        st.session_state.get("optimized_resume_path")
        and os.path.exists(st.session_state["optimized_resume_path"])
    )

    tab_a, tab_b = st.tabs(["🔄 转当前优化简历", "📤 上传其他 Word 文件"])

    with tab_a:
        if has_optimized:
            st.info(f"📄 当前优化简历：**{st.session_state['optimized_resume_name']}**")
            col_a1, _, col_a2 = st.columns([2, 1, 3])
            with col_a1:
                if st.button("🔄 一键转为 PDF", type="primary",
                             use_container_width=True,
                             key="btn_convert_optimized"):
                    with st.spinner("⏳ 正在转换 Word → PDF..."):
                        try:
                            pdf_path, method = convert_word_to_pdf(
                                st.session_state["optimized_resume_path"]
                            )
                            st.session_state["generated_pdf_path"] = pdf_path
                            st.session_state["generated_pdf_name"] = os.path.basename(pdf_path)
                            st.session_state["step_completed"]["step6"] = True
                            st.rerun()
                        except Exception as e:
                            st.error(f"❌ PDF 转换失败：{e}")

            # 显示 PDF 下载按钮
            pdf_path = st.session_state.get("generated_pdf_path")
            if pdf_path and os.path.exists(pdf_path):
                st.success(f"✅ PDF 已生成！（转换引擎：{'LibreOffice 高质量转换' if 'libreoffice' in str(pdf_path) else '本地引擎转换'}）")
                with open(pdf_path, "rb") as f:
                    pdf_data = f.read()
                st.download_button(
                    label=f"📥 下载 PDF ({st.session_state['generated_pdf_name']})",
                    data=pdf_data,
                    file_name=st.session_state["generated_pdf_name"],
                    mime="application/pdf",
                    use_container_width=True,
                    key="download_pdf_optimized",
                )
                st.caption(f"文件大小：{len(pdf_data) / 1024:.0f} KB")
        else:
            st.warning("⚠️ 请先完成步骤 5（改写简历），或切换到「上传其他 Word 文件」标签页")

    with tab_b:
        uploaded_docx = st.file_uploader(
            "选择 Word 文件 (.docx)",
            type=["docx"],
            help="上传需要转为 PDF 的 Word 简历文件",
            key="upload_docx_for_pdf",
        )

        if uploaded_docx:
            # 用原文件名展示
            st.info(f"📄 已选择：**{uploaded_docx.name}**")
            col_b1, _, col_b2 = st.columns([2, 1, 3])
            with col_b1:
                if st.button("🔄 上传并转为 PDF", type="primary",
                             use_container_width=True,
                             key="btn_convert_upload"):
                    with st.spinner("⏳ 正在转换 Word → PDF..."):
                        try:
                            # 用原文件名保存到系统临时目录（而不是 NamedTemporaryFile 随机名）
                            original_name = uploaded_docx.name
                            base_name = os.path.splitext(original_name)[0]
                            # 放在 /tmp/<原文件名>.docx，转换后 PDF 也会用同 base_name
                            safe_name = re.sub(r'[^\w\u4e00-\u9fff\-_.]', '_', base_name)
                            tmp_dir = tempfile.gettempdir()
                            docx_path = os.path.join(tmp_dir, f"{safe_name}.docx")
                            with open(docx_path, "wb") as f:
                                f.write(uploaded_docx.getvalue())

                            pdf_path, method = convert_word_to_pdf(docx_path)
                            # 清理上传的中间 docx
                            try:
                                os.unlink(docx_path)
                            except Exception:
                                pass

                            st.session_state["generated_pdf_path"] = pdf_path
                            st.session_state["generated_pdf_name"] = os.path.basename(pdf_path)
                            st.session_state["step_completed"]["step6"] = True
                            st.rerun()
                        except Exception as e:
                            st.error(f"❌ PDF 转换失败：{e}")

            # 显示 PDF 下载按钮
            pdf_path = st.session_state.get("generated_pdf_path")
            if pdf_path and os.path.exists(pdf_path):
                st.success(f"✅ PDF 已生成！")
                with open(pdf_path, "rb") as f:
                    pdf_data = f.read()
                st.download_button(
                    label=f"📥 下载 PDF ({st.session_state['generated_pdf_name']})",
                    data=pdf_data,
                    file_name=st.session_state["generated_pdf_name"],
                    mime="application/pdf",
                    use_container_width=True,
                    key="download_pdf_upload",
                )
                st.caption(f"文件大小：{len(pdf_data) / 1024:.0f} KB")

    # ========================================================================
    # 页脚
    # ========================================================================
    st.markdown(
        """
        <div style="text-align: center; color: #999; padding: 30px 0; font-size: 13px;">
        📄 智能简历匹配与改写工具（本地版） | 纯本地处理 | 无需联网 | 无需 API Key
        </div>
        """,
        unsafe_allow_html=True,
    )


# ============================================================================
# 主入口
# ============================================================================

def main():
    """主函数"""
    render_sidebar()
    render_main_area()


if __name__ == "__main__":
    main()
